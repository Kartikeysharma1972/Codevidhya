print("[DEBUG] main.py loading...")

# Fix for Python 3.14 Windows asyncio issues
import asyncio
import sys
if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

from fastapi import FastAPI, HTTPException, UploadFile, File, Request
import io
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from openai import OpenAI
from dotenv import load_dotenv
import os
import json
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from database import db
from mcp_tools import get_chat_history, check_usage_limit, increment_usage
from auth import register_user, login_user, verify_token, logout_user

# ─── VOCABULARY MASTERY & READING COMPREHENSION TOOL IMPORTS ─────
import re as _re_mod
import uuid as _uuid_mod

# Vocabulary Mastery supporting modules (renamed copies)
from vocab_nlp import (
    get_grade_prompt_context as vocab_get_grade_prompt_context,
    analyze_text_grade as vocab_analyze_text_grade,
    GRADE_PROFILES as VOCAB_GRADE_PROFILES,
    get_word_count as vocab_get_word_count,
)
from vocab_rag import rag_retriever as vocab_rag_retriever
from vocab_db import (
    init_db as vocab_init_db,
    create_session as vocab_create_session,
    save_worksheet as vocab_save_worksheet,
    get_session_history as vocab_get_session_history,
    get_all_worksheets as vocab_get_all_worksheets,
    save_rag_document as vocab_save_rag_document,
)
from tool_security import (
    assert_public_url, read_upload_capped, client_ip,
    generate_limiter, extract_limiter, upload_limiter,
)

# Reading Comprehension supporting modules (renamed copies)
from reading_nlp import (
    get_grade_prompt_context as reading_get_grade_prompt_context,
    analyze_text_grade as reading_analyze_text_grade,
    GRADE_PROFILES as READING_GRADE_PROFILES,
    get_reading_counts,
)
from reading_db import (
    init_db as reading_init_db,
    create_session as reading_create_session,
    save_comprehension as reading_save_comprehension,
    get_session_history as reading_get_session_history,
    get_all_comprehensions as reading_get_all_comprehensions,
    save_rag_document as reading_save_rag_document,
    get_all_rag_documents as reading_get_all_rag_documents,
)

# Initialize tool databases at import time
try:
    vocab_init_db()
except Exception as _e:
    print(f"[startup] Vocab DB init error: {_e}")
try:
    reading_init_db()
except Exception as _e:
    print(f"[startup] Reading DB init error: {_e}")
try:
    vocab_rag_retriever.build_index()
except Exception as _e:
    print(f"[startup] Vocab RAG build error: {_e}")

print("[DEBUG] All imports complete")
load_dotenv()
print("[DEBUG] Creating FastAPI app...")

app = FastAPI(title="ClassroomAI API")

# Import RAG modules (available after pip install)
try:
    from rag_agent import RAGAgent
    from code_analyzer import CodeAnalyzer
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False

_CORS_ORIGINS = os.getenv("CORS_ORIGINS", "http://localhost:5176,http://localhost:5173,http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=_CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OPENAI_API_KEY = os.getenv("GROQ_API_KEY", "").strip() or "missing-set-GROQ_API_KEY-in-env"
OPENAI_MODEL   = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
client = OpenAI(api_key=OPENAI_API_KEY, base_url="https://api.groq.com/openai/v1")

# ─── AUTH MODELS & ENDPOINTS ─────────────────────────

class RegisterRequest(BaseModel):
    name: str
    email: str
    password: str
    school_name: str = ""

class LoginRequest(BaseModel):
    email: str
    password: str

def get_current_user(request: Request):
    auth = request.headers.get("Authorization", "")
    token = auth.replace("Bearer ", "") if auth.startswith("Bearer ") else ""
    user = verify_token(token)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return user

@app.post("/api/auth/register")
async def api_register(req: RegisterRequest):
    result = register_user(req.name, req.email, req.password, req.school_name)
    if not result["success"]:
        raise HTTPException(status_code=400, detail=result["error"])
    return result

@app.post("/api/auth/login")
async def api_login(req: LoginRequest):
    result = login_user(req.email, req.password)
    if not result["success"]:
        raise HTTPException(status_code=401, detail=result["error"])
    return result

@app.post("/api/auth/logout")
async def api_logout(request: Request):
    auth = request.headers.get("Authorization", "")
    token = auth.replace("Bearer ", "") if auth.startswith("Bearer ") else ""
    logout_user(token)
    return {"success": True}

@app.get("/api/auth/me")
async def api_me(request: Request):
    user = get_current_user(request)
    return {"success": True, "user": user}

# ─── MODELS ───────────────────────────────────────────

class WorksheetRequest(BaseModel):
    topic: str
    grade_level: str
    subject: str = ""
    worksheet_type: str = "mixed"
    num_questions: int = 10
    differentiation_level: str = "grade-level"
    blooms_level: str = "mixed"
    include_word_bank: bool = False
    additional_instructions: str = ""
    source_material: str = ""
    topic_description: str = ""
    topic_track: str = "core"

class LessonPlanRequest(BaseModel):
    topic: str
    grade_level: str
    subject: str
    duration: str = "45 minutes"
    objectives: str = ""
    standards: str = ""
    class_type: str = "in-person"
    learning_style: str = "mixed"
    student_needs: str = "general"
    tech_integration: str = "low"
    include_topic_overview: bool = True
    additional_notes: str = ""
    source_material: str = ""
    # Optional NCERT/CBSE chapter description for the selected topic.
    topic_description: str = ""
    # "core" (CBSE TOC topic) or "miscellaneous" (curated extra topic).
    topic_track: str = "core"

class LessonPlanEnrichRequest(BaseModel):
    # The existing lesson plan body (so the model can extend it coherently).
    existing_plan: str
    topic: str
    grade_level: str
    subject: str
    duration: str = "45 minutes"
    topic_description: str = ""
    # One of: "more_activities", "more_examples", "more_topics"
    action: str = "more_activities"

class MCAssessmentRequest(BaseModel):
    topic: str
    grade_level: str
    subject: str = ""
    num_questions: int = 10
    difficulty: str = "medium"
    blooms_level: str = "mixed"
    question_format: str = "pure_mc"
    include_explanations: bool = True
    standards: str = ""
    additional_instructions: str = ""
    source_material: str = ""

# ─── RAG ENDPOINT MODELS ──────────────────────────────

class CodeExplainRequest(BaseModel):
    code: str
    language: str
    grade: str

class DebugRequest(BaseModel):
    code: str
    language: str
    error: str

class ImproveRequest(BaseModel):
    code: str
    language: str
    focus: str = "best_practices"

class AnalyzeRequest(BaseModel):
    code: str
    language: str = "python"

class PatternRequest(BaseModel):
    pattern: str
    language: str
    grade: str

# ─── CHAT HISTORY & USAGE MODELS ──────────────────────

class ChatHistoryRequest(BaseModel):
    teacher_id: str

class UsageCheckRequest(BaseModel):
    teacher_id: str
    tool_name: str

class UsageIncrementRequest(BaseModel):
    teacher_id: str
    tool_name: str

class SaveChatRequest(BaseModel):
    teacher_id: str
    tool_name: str
    topic: str
    grade_level: str
    subject: str
    request_data: dict
    response_preview: str
    response_content: str = None  # Full response content

# ─── HELPERS ──────────────────────────────────────────

def parse_grade_number(grade_level: str):
    """Map a free-text grade label to a numeric level.
    Kindergarten/pre-primary -> 0, Grade 1..12 -> 1..12, College/UG -> 13.
    Returns None if it cannot be determined."""
    g = (grade_level or "").lower()
    if any(x in g for x in ["kindergarten", "pre-primary", "pre primary", "nursery", "kg", "lkg", "ukg", "k-"]):
        return 0
    if any(x in g for x in ["college", "university", "undergrad", "graduate", "ug ", "btech", "b.tech"]):
        return 13
    m = _re_mod.search(r"\b(?:grade|class|std|standard)\s*([0-9]{1,2})\b", g)
    if not m:
        m = _re_mod.search(r"\b([0-9]{1,2})\s*(?:st|nd|rd|th)\b", g)
    if m:
        n = int(m.group(1))
        if 1 <= n <= 12:
            return n
    if "high school" in g:
        return 10
    if "middle" in g:
        return 7
    return None


# Per-grade cognitive/intelligence calibration. Each entry tells the model the EXACT
# age, vocabulary ceiling, sentence length, cognitive (Bloom's) ceiling, reasoning depth,
# expected answer length, and the kind of examples that fit a student of that grade.
# This is what makes a Grade 1 output genuinely different from a Grade 10 output.
_GRADE_INTELLIGENCE = {
    0:  ("age 4-5 (Kindergarten)", "only spoken sight words a 4-5 year old knows", "3-5 words",
         "Remember & basic Understand only (name, point, match, count to 20)",
         "no abstract reasoning; one idea at a time; recognise and repeat",
         "one word or one short spoken sentence",
         "toys, animals, food, family, colours, shapes — things a child can touch or see"),
    1:  ("age 6 (Grade 1)", "basic sight words and CVC words a 6 year old reads", "4-6 words",
         "Remember & Understand (name, list, match, count, identify)",
         "single-step reasoning; concrete only; pictures help",
         "1 short sentence or a single number",
         "everyday home/classroom objects, counting things, simple picture stories"),
    2:  ("age 7 (Grade 2)", "simple everyday words a 7 year old reads", "5-8 words",
         "Remember & Understand, very simple Apply (single-step add/subtract, sort, order)",
         "one or two concrete steps; no abstraction",
         "1-2 short sentences",
         "money, time, simple shopping, family, short stories with a clear moral"),
    3:  ("age 8 (Grade 3)", "common everyday vocabulary a 8 year old knows", "8-10 words",
         "Understand & Apply; first simple 'why' questions",
         "two-step concrete reasoning; cause and effect in plain terms",
         "2-3 sentences or short working",
         "real situations: markets, gardens, weather, simple maps, classroom surveys"),
    4:  ("age 9 (Grade 4)", "elementary vocabulary; define any new term immediately", "10-12 words",
         "Apply & begin Analyse (compare, group, explain reasons)",
         "multi-step concrete reasoning; simple patterns and rules",
         "3-4 sentences or a few lines of working",
         "daily-life word problems, simple experiments, comparisons, short paragraphs"),
    5:  ("age 10 (Grade 5)", "elementary vocabulary with a few subject terms (defined)", "10-14 words",
         "Apply & Analyse; introduce light abstraction",
         "multi-step reasoning; generalise simple rules; justify briefly",
         "4-5 sentences or structured working",
         "real-world projects, data tables, fractions in cooking/money, cause-effect in stories"),
    6:  ("age 11 (Grade 6)", "middle-school vocabulary; introduce subject terms with brief definitions", "moderate (12-16 words)",
         "Understand, Apply & Analyse; early abstract reasoning",
         "multi-step structured reasoning; handle variables and simple models",
         "5-8 sentences or full step-by-step working",
         "real scenarios, simple algebra/geometry, science with given data, source-based questions"),
    7:  ("age 12 (Grade 7)", "middle-school vocabulary; subject terms used with short context", "moderate (12-18 words)",
         "Apply & Analyse; form simple hypotheses",
         "multi-step reasoning across 2 ideas; compare and infer",
         "6-10 sentences or detailed working",
         "applied problems, experiments, map/source analysis, multi-step word problems"),
    8:  ("age 13 (Grade 8)", "subject-specific vocabulary used naturally", "varied, can be complex",
         "Analyse & begin Evaluate; pre-algebra/abstract thinking",
         "integrate 2-3 concepts; reason about relationships and patterns",
         "detailed multi-paragraph or full derivations",
         "case studies, data interpretation, derivations, real-world modelling"),
    9:  ("age 14 (Grade 9)", "discipline-specific terminology, lightly defined", "complex sentences expected",
         "Analyse, Evaluate & some Create; formal abstract reasoning",
         "multi-concept integration; proofs, derivations, justified arguments",
         "full board-style answers with steps and justification",
         "board-exam scenarios, derivations, HOTS, critical analysis of texts/data"),
    10: ("age 15 (Grade 10 — board year)", "discipline-specific terminology used freely", "complex, academic",
         "Analyse, Evaluate & Create; full higher-order thinking",
         "rigorous multi-concept integration; complete proofs and critical evaluation",
         "complete board-quality answers with reasoning and justification",
         "CBSE board / exemplar-level problems, HOTS, multi-step proofs, source/data analysis"),
    11: ("age 16 (Grade 11)", "advanced academic and technical terminology", "complex, theoretical",
         "Evaluate & Create; theoretical depth and synthesis",
         "abstract, theory-driven reasoning; connect multiple concepts and derive results",
         "extended analytical answers with rigorous justification",
         "competitive-exam (JEE/NEET foundation) style, theoretical derivations, research-style questions"),
    12: ("age 17 (Grade 12)", "advanced academic and technical terminology", "complex, theoretical",
         "Evaluate & Create; pre-college rigor",
         "high abstraction; multi-step derivations, critical synthesis, original argument",
         "extended, rigorous, exam-quality answers",
         "JEE/NEET/board-advanced problems, in-depth analysis, applied theory"),
    13: ("undergraduate (College)", "specialised academic and professional terminology", "scholarly",
         "Evaluate & Create at university level",
         "independent, research-grade reasoning; critique and construct arguments",
         "essay/derivation-length scholarly answers",
         "research-oriented, real-world professional and theoretical problems"),
}


def get_grade_language_profile(grade_level: str) -> str:
    """Return a precise cognitive/intelligence calibration block for the given grade.
    This is what differentiates a Grade 1 output from a Grade 10 output — vocabulary,
    sentence length, the Bloom's ceiling, reasoning depth, and expected answer length
    are all pinned to the exact grade so the model never produces above/below-level work."""
    n = parse_grade_number(grade_level)
    if n is None:
        return "GRADE INTELLIGENCE LEVEL: Use clear, age-appropriate language and reasoning depth for the specified grade level."
    age, vocab, sentence, cognition, reasoning, answer_len, examples = _GRADE_INTELLIGENCE[n]
    return (
        f"GRADE INTELLIGENCE LEVEL — CRITICAL, calibrate EVERYTHING to {grade_level} ({age}):\n"
        f"  • VOCABULARY: use {vocab}. Never use words above this level; if a new term is unavoidable, define it in simpler words.\n"
        f"  • SENTENCE LENGTH: {sentence}.\n"
        f"  • COGNITIVE CEILING (Bloom's): {cognition}. Do NOT exceed this — a {grade_level} student cannot handle higher-grade cognition.\n"
        f"  • REASONING DEPTH: {reasoning}.\n"
        f"  • EXPECTED ANSWER LENGTH: {answer_len}.\n"
        f"  • EXAMPLES MUST USE: {examples}.\n"
        f"  The difficulty, wording, and depth must be UNMISTAKABLY {grade_level} — clearly simpler than the grade above and clearly more advanced than the grade below."
    )


def get_curriculum_guardrails(grade_level: str, subject: str = "", topic: str = "",
                              topic_description: str = "", strict_syllabus: bool = True) -> str:
    """Shared, non-negotiable rules injected into every generation prompt:
    (1) lock content to the chosen grade's level, (2) forbid repeated/duplicate items,
    (3) make the model reason about what/how/why before writing. Centralised so every
    tool (worksheet, lesson plan, question paper, quiz, activities, assessment) behaves the same.

    strict_syllabus=True (CBSE/NCERT core track): stay strictly inside the grade's textbook syllabus.
    strict_syllabus=False (Miscellaneous / Advanced track): may go beyond the textbook for enrichment,
    but difficulty and cognition MUST stay calibrated to the chosen grade (never higher/lower)."""
    subj = f" ({subject})" if subject else ""
    if strict_syllabus:
        rule1 = (
            f"1. SYLLABUS LOCK: Produce content ONLY at the {grade_level} level for '{topic}'{subj}. "
            f"Stay strictly within what a {grade_level} student studies for this topic in their syllabus. "
            f"Do NOT borrow concepts, vocabulary, examples, formulae, or sub-topics that belong to a higher OR a lower grade. "
            f"If something is normally taught in a different class, leave it out.\n"
        )
    else:
        rule1 = (
            f"1. GRADE-LEVEL LOCK: Content is for '{topic}'{subj} at the {grade_level} level. "
            f"You MAY go beyond the standard textbook for enrichment, BUT the difficulty, vocabulary, and reasoning "
            f"demand MUST stay calibrated to a {grade_level} student — never use the complexity or techniques of a "
            f"higher or lower grade. Keep it challenging-but-reachable for this exact grade.\n"
        )
    block = (
        "═══ NON-NEGOTIABLE RULES (apply to the ENTIRE output) ═══\n"
        + rule1 +
        "2. ZERO REPETITION: Every question / item / example MUST be unique. No two items may test the same fact, "
        "reuse the same numbers or scenario, or be a reworded version of another. Vary the sub-concept, the context, "
        "the numbers, and the phrasing across all items. Before finalising, re-read everything and replace any "
        "duplicate or near-duplicate so the same question never appears twice.\n"
        "3. THINK BEFORE WRITING (silently): For each item first decide WHAT concept it tests, HOW a student of this exact "
        "grade would solve it, and WHY it suits this grade — then write a clean, correct item. "
        "Output ONLY the final content, never your reasoning.\n"
    )
    if topic_description:
        block += f"4. CURRICULUM ANCHOR: Keep every item aligned to this scope — {topic_description}\n"
    block += "═════════════════════════════════════════════════════════\n"
    return block


def _dedupe_questions(questions: list) -> list:
    """Remove exact / near-exact duplicate questions (normalised text match).
    Guarantees the same question never appears twice in a generated paper/quiz."""
    seen = set()
    out = []
    for q in questions or []:
        text = (q.get("question") if isinstance(q, dict) else None) or ""
        key = _re_mod.sub(r"[^a-z0-9]", "", text.lower())
        if key and key in seen:
            continue
        if key:
            seen.add(key)
        out.append(q)
    return out


def call_openai(system_prompt: str, user_prompt: str, max_tokens: int = 4096, temperature: float = 0.5) -> str:
    try:
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI error: {str(e)}")

# ─── ROUTES ───────────────────────────────────────────

@app.post("/api/worksheet")
def generate_worksheet(req: WorksheetRequest):
    if not req.topic.strip():
        raise HTTPException(status_code=400, detail="Topic is required.")

    type_map = {
        "fill_blank":      "fill-in-the-blank questions (use ________ for blanks)",
        "multiple_choice": "multiple choice questions with 4 options each (A, B, C, D)",
        "open_ended":      "open-ended short answer questions requiring critical thinking",
        "mixed":           "a balanced mix of fill-in-the-blank, multiple choice, and open-ended questions",
        "qa":              "question and answer pairs formatted as 'Q: [question]' then 'A: [answer]' on separate lines",
    }

    diff_map = {
        "grade-level": "exactly at grade level using standard curriculum vocabulary",
        "beginner":    "simplified and scaffolded for struggling or below-grade-level students using shorter sentences and simpler vocabulary",
        "advanced":    "challenging for above-grade-level students requiring deeper analysis and extended thinking",
        "mixed":       "differentiated with a range of difficulty from foundational to higher-order thinking to support all learners",
    }

    blooms_map = {
        "remember":   "remembering and recalling facts (define, list, identify, name) — Bloom's Level 1",
        "understand": "understanding and explaining concepts (describe, explain, summarize) — Bloom's Level 2",
        "apply":      "applying knowledge to new situations (use, solve, demonstrate, calculate) — Bloom's Level 3",
        "analyze":    "analyzing and breaking down information (compare, contrast, differentiate) — Bloom's Level 4",
        "evaluate":   "evaluating and making judgments (assess, critique, justify, argue) — Bloom's Level 5",
        "create":     "creating and synthesizing new ideas (design, compose, construct, formulate) — Bloom's Level 6",
        "mixed":      "a balanced range across all Bloom's Taxonomy levels from recall to higher-order thinking",
    }

    q_type  = type_map.get(req.worksheet_type, type_map["mixed"])
    diff    = diff_map.get(req.differentiation_level, diff_map["grade-level"])
    blooms  = blooms_map.get(req.blooms_level, blooms_map["mixed"])

    word_bank_note = ""
    if req.include_word_bank and req.worksheet_type in ("fill_blank", "mixed"):
        word_bank_note = (
            "Include a WORD BANK box near the top of the worksheet listing all the missing words "
            "students must use, arranged in a bordered box with words separated by commas."
        )

    material_note = ""
    if req.source_material.strip():
        material_note = (
            f"\n\nTEACHER-UPLOADED SOURCE MATERIAL (base your questions primarily on this content):\n"
            f"---\n{req.source_material[:5000]}\n---\n"
        )

    lang_profile = get_grade_language_profile(req.grade_level)

    # Get subject-specific question intelligence
    subj_lower = (req.subject or '').lower()
    subject_intelligence = ""
    if 'math' in subj_lower:
        subject_intelligence = (
            "MATH WORKSHEET RULES:\n"
            "- Every question must require CALCULATION or MATHEMATICAL WORKING\n"
            "- Include specific numerical values in every problem\n"
            "- Use word problems with real-world scenarios (not just 'Solve 2+3')\n"
            "- For higher grades: include proofs, constructions, derivations\n"
            "- NEVER write 'Define X' or 'What is X' — always ask 'Find', 'Calculate', 'Prove', 'Solve', 'Construct'\n"
        )
    elif any(s in subj_lower for s in ['science', 'physics', 'chemistry', 'biology']):
        subject_intelligence = (
            "SCIENCE WORKSHEET RULES:\n"
            "- Include numerical problems with given data and formulas\n"
            "- Ask diagram-based questions (draw and label, identify parts)\n"
            "- Include experiment-based questions (procedure, observation, conclusion)\n"
            "- Use real-world application scenarios, not just textbook recall\n"
            "- NEVER write generic 'Define X' or 'Explain X' — ask specific mechanism/reasoning questions\n"
        )
    elif any(s in subj_lower for s in ['history', 'social', 'geography', 'civics']):
        subject_intelligence = (
            "SOCIAL SCIENCE WORKSHEET RULES:\n"
            "- Include source-based questions with excerpts\n"
            "- Ask map-based questions for geography\n"
            "- Focus on cause-effect, analysis, and evaluation — not just dates/names\n"
            "- Include case study and scenario-based questions\n"
        )
    elif any(s in subj_lower for s in ['artificial intelligence', 'computer', 'information technology', 'coding', 'programming']) or subj_lower.strip() in ('ai', 'it', 'cs'):
        subject_intelligence = (
            "AI / COMPUTER SCIENCE / IT WORKSHEET RULES:\n"
            "- Questions MUST be about technology concepts, NOT mathematical computation\n"
            "- DO NOT generate math problems with tech-related object names (robots, AI, etc.)\n"
            "- Include scenario-based questions about real-world AI/tech applications\n"
            "- Ask about how AI is used in: healthcare, agriculture, education, transport, entertainment\n"
            "- Include questions about ethics, data privacy, bias in AI, responsible AI use\n"
            "- Reference real tools: Google AI, ChatGPT, Python, Scratch, Alexa/Siri, self-driving cars\n"
            "- Test understanding of concepts: machine learning, neural networks, NLP, computer vision\n"
            "- For MCQ: options should be technology concepts, NOT numbers\n"
            "- NEVER write math computation questions disguised as AI/tech questions\n"
        )
    elif 'english' in subj_lower:
        subject_intelligence = (
            "ENGLISH WORKSHEET RULES:\n"
            "- Include passage-based comprehension with inference questions\n"
            "- Grammar questions should test APPLICATION, not just identification\n"
            "- Include creative writing prompts with clear guidelines\n"
            "- Test vocabulary in context, not isolated definitions\n"
            "- For literature: character analysis, theme exploration, critical evaluation\n"
        )
    elif 'hindi' in subj_lower:
        subject_intelligence = (
            "HINDI WORKSHEET RULES:\n"
            "- Include passage-based comprehension (apathit gadyansh)\n"
            "- Grammar: sandhi-viched, samas, alankar, muhavare in context\n"
            "- Creative writing: patra, nibandh, kahani with clear instructions\n"
            "- Literature: kavya-bodh, character analysis, theme discussion\n"
        )

    # Build subject-aware system prompt
    is_tech = any(s in subj_lower for s in ['artificial intelligence', 'computer', 'information technology', 'coding', 'programming']) or subj_lower.strip() in ('ai', 'it', 'cs')

    if is_tech:
        system_intro = (
            f"You are a senior {req.subject} teacher with expertise in technology education. "
            f"Every question you write is about REAL technology concepts, applications, and scenarios. "
            f"You NEVER write math computation questions disguised as tech questions. "
            f"You NEVER just replace object names (robots instead of apples) in math problems. "
            f"Your questions test genuine understanding of HOW technology works and WHERE it is applied. "
        )
    else:
        system_intro = (
            "You are a senior CBSE/NCERT teacher with 20 years of experience creating worksheets. "
            "Every question you write is SPECIFIC, requires genuine thinking, and matches the quality "
            "of actual CBSE board exam questions. You NEVER write lazy questions like 'Define X' or 'What is X'. "
            "Every question must require analysis, reasoning, or creative application. "
        )

    system_prompt = (
        system_intro
        + f"{lang_profile} "
        + ("Derive ALL questions from the provided source material. " if req.source_material.strip() else "")
        + "Format output with CLEAN MARKDOWN for readability:\n"
        "- Use # for main title, ## for section headers, ### for sub-sections\n"
        "- Use **bold** for important terms and labels\n"
        "- Use numbered lists (1. 2. 3.) for questions\n"
        "- Use bullet points (- ) for lists\n"
        "- Use --- for section dividers\n"
    )

    guardrails = get_curriculum_guardrails(req.grade_level, req.subject, req.topic, req.topic_description or "", strict_syllabus=(req.topic_track == "core"))

    user_prompt = (
        f"Create a PRINT-READY worksheet for {req.grade_level} students.\n\n"
        f"TOPIC: {req.topic}\n"
        f"{'SUBJECT: ' + req.subject if req.subject else ''}\n"
        f"QUESTION TYPE: {q_type}\n"
        f"NUMBER OF QUESTIONS: {req.num_questions}\n"
        f"DIFFERENTIATION LEVEL: {diff}\n"
        f"BLOOM'S TAXONOMY FOCUS: {blooms}\n"
        f"{word_bank_note}\n"
        f"{'ADDITIONAL INSTRUCTIONS: ' + req.additional_instructions if req.additional_instructions else ''}\n"
        f"{material_note}\n\n"
        f"{guardrails}\n"
        f"{subject_intelligence}\n"
        "QUESTION QUALITY RULES (CRITICAL):\n"
        "- EVERY question must be a complete, well-formed sentence as it would appear on a printed exam\n"
        "- Include specific values, measurements, names, dates, or scenarios in EVERY question\n"
        "- For numerical subjects: at least 50% questions must require step-by-step calculation\n"
        "- FORBIDDEN: 'Define X', 'What is X', 'List the features of X', 'Explain X in brief'\n"
        "- Each question must need at least 2-3 lines of working/writing to answer\n\n"
        "FORMAT:\n"
        "- Title in ALL CAPS\n"
        "- Name / Date / Class lines\n"
        f"{'- Word Bank box listing all fill-in words' if req.include_word_bank else ''}\n"
        "- Brief student instructions\n"
        "- Questions numbered 1, 2, 3...\n"
        "- For MCQ: 4 plausible options (A, B, C, D) — no obviously wrong choices\n"
        "- For fill-in-blank: use ________ for blanks\n"
        "- For open-ended: include point values and expected answer length\n\n"
        "## Answer Key\n"
        "- Each answer numbered to match\n"
        "- Include full solution steps for calculation questions\n"
        "- Include 1-2 sentence explanation for each answer\n\n"
        "Use proper markdown formatting: # for title, ## for sections, **bold** for key terms, numbered lists for questions, - for bullets."
    )

    # Scale token budget with question count so answer keys never truncate.
    # Each question + answer ~ 350-450 tokens; add 600 for title/instructions/word bank.
    qcount = max(1, int(req.num_questions or 10))
    worksheet_tokens = min(8000, max(4096, 600 + qcount * 420))
    result = call_openai(system_prompt, user_prompt, max_tokens=worksheet_tokens)
    return {"result": result, "tool": "worksheet"}


@app.post("/api/upload-material")
async def upload_material(file: UploadFile = File(...)):
    """Extract text from teacher-uploaded material (PDF, DOCX, TXT, MD)"""
    content = await file.read()
    filename = (file.filename or "").lower()
    try:
        if filename.endswith(".pdf"):
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                text = "\n".join(p.extract_text() or "" for p in reader.pages)
            except ImportError:
                raise HTTPException(status_code=400, detail="PDF support not installed. Upload a .txt file instead.")
        elif filename.endswith(".docx"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                raise HTTPException(status_code=400, detail="DOCX support not installed. Upload a .txt file instead.")
        else:
            text = content.decode("utf-8", errors="ignore")
        text = text.strip()[:8000]
        if not text:
            raise HTTPException(status_code=400, detail="Could not extract text from file.")
        return {"text": text, "chars": len(text), "filename": file.filename}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not process file: {str(e)}")


@app.post("/api/lesson-plan")
def generate_lesson_plan(req: LessonPlanRequest):
    if not req.topic.strip():
        raise HTTPException(status_code=400, detail="Topic is required.")

    lang_profile = get_grade_language_profile(req.grade_level)

    style_map = {
        "mixed":       "visual + auditory + kinesthetic combined",
        "visual":      "diagrams, charts, graphic organizers, color-coding, videos",
        "auditory":    "discussion, verbal explanations, think-pair-share, reading aloud",
        "kinesthetic": "movement activities, manipulatives, experiments, role-play",
    }

    needs_map = {
        "general":       "a general education classroom with diverse learners",
        "ell":           "English Language Learners — sentence frames, visual vocabulary, bilingual supports",
        "special_needs": "students with diverse learning needs — modifications, accommodations, chunked tasks",
        "gifted":        "gifted and talented — enrichment tasks, higher-order questioning, independent projects",
    }

    tech_map = {
        "low":         "low-tech: paper, pencil, physical manipulatives, whiteboards",
        "digital":     "digital: presentations, educational websites, online quizzes, Google Docs/Slides",
        "interactive": "interactive/blended: interactive whiteboard, simulations, collaborative platforms, apps",
    }

    class_map = {
        "in-person": "traditional in-person classroom",
        "remote":    "fully remote/virtual — video call strategies, breakout rooms, digital submission",
        "hybrid":    "hybrid with both in-person and remote students simultaneously",
    }

    lesson_material_note = ""
    if req.source_material.strip():
        lesson_material_note = (
            f"\n\nTEACHER-UPLOADED SOURCE MATERIAL (align the lesson to this content):\n"
            f"---\n{req.source_material[:5000]}\n---\n"
        )

    chapter_context = ""
    if req.topic_description.strip():
        track_label = "OFFICIAL CBSE/NCERT CHAPTER" if req.topic_track == "core" else "ENRICHMENT TOPIC"
        chapter_context = (
            f"\n{track_label} REFERENCE:\n\"{req.topic_description.strip()}\"\n"
        )

    additional_notes_instruction = ""
    if req.additional_notes and req.additional_notes.strip():
        additional_notes_instruction = (
            f"\n\nTEACHER'S SPECIAL REQUEST (MUST follow this):\n"
            f"\"{req.additional_notes.strip()}\"\n"
            f"Integrate this request throughout the lesson plan. If the teacher asks for "
            f"example questions at each subtopic, include worked examples under every subtopic. "
            f"If the teacher asks for specific activities, include those activities.\n"
        )

    additional_notes_section = ""
    if req.additional_notes and req.additional_notes.strip():
        additional_notes_section = (
            f"\nTEACHER'S SPECIAL REQUEST: {req.additional_notes.strip()}\n"
            f"You MUST follow this request and integrate it into the lesson plan.\n"
        )

    # Parse class duration into total minutes for accurate time-boxing across the 9 sections.
    duration_total = 45
    try:
        if "hour" in req.duration.lower():
            duration_total = int(req.duration.split()[0]) * 60
        elif "block" in req.duration.lower():
            duration_total = 90
        else:
            duration_total = int(req.duration.split()[0])
    except Exception:
        duration_total = 45

    # Suggested time split across the 9 sections (sums to duration_total).
    weights = [0.05, 0.05, 0.10, 0.20, 0.15, 0.15, 0.10, 0.10, 0.10]
    times = [max(2, round(duration_total * w)) for w in weights]
    # Make times sum exactly to duration_total
    drift = duration_total - sum(times)
    times[3] += drift  # absorb drift into Direct Instruction
    section_names = [
        "Objective & Competency Anchor",
        "Prior Knowledge Activation",
        "Hook & Real-World Connect",
        "Direct Instruction",
        "Guided Practice",
        "Independent Practice",
        "HOTS Task",
        "Formative Assessment",
        "Closure & Home Connect",
    ]
    time_table = "\n".join(f"  {i+1}. {section_names[i]} — {times[i]} min" for i in range(9))

    standards_block = ""
    if req.standards and req.standards.strip():
        standards_block = f"\nStandards / Curriculum: {req.standards.strip()}\n"

    system_prompt = (
        "You are an expert teacher and instructional designer creating a classroom-ready 9-section "
        "lesson plan that another teacher can walk in and teach from. You always include ACTUAL solved "
        "examples with real numbers and step-by-step working. You never write generic descriptions.\n\n"
        "OUTPUT FORMAT (STRICT):\n"
        "- Respond in CLEAN MARKDOWN ONLY. No JSON, no HTML, no whole-output code fences, no [TO BE FILLED] placeholders.\n"
        "- Start with a single # title line, then 9 ## sections in the EXACT order specified.\n"
        "- Every section header MUST include its time allocation in parentheses, e.g. '## 1. Objective & Competency Anchor (3 min)'.\n"
        "- Sum of the 9 section times MUST equal the total class duration. Never overflow.\n"
        "- If a section is genuinely not applicable (e.g. no prior knowledge to activate for an intro topic, or no HOTS at "
        "lower grades), output the section header with the note '_Flagged: this section may be skipped — <reason>._' "
        "but DO still include the section header so the teacher sees the structure.\n"
        "- Use **bold** for key terms, numbered lists for steps, - for bullet points, --- between sections.\n"
        "- For math: plain text (e.g. x^2 + 2x + 1). NO LaTeX $$ or \\(...\\).\n"
        "- Replace every [square-bracket placeholder] with concrete content.\n\n"
        f"{lang_profile}"
    )

    user_prompt = (
        f"Write a 9-section, time-boxed lesson plan.\n\n"
        f"Topic: {req.topic}\n"
        f"Subject: {req.subject}\n"
        f"Grade: {req.grade_level}\n"
        f"Total Class Duration: {req.duration} ({duration_total} minutes)\n"
        f"{standards_block}"
        f"{chapter_context}"
        f"{additional_notes_section}"
        f"{lesson_material_note}\n\n"
        f"{get_curriculum_guardrails(req.grade_level, req.subject, req.topic, req.topic_description or '', strict_syllabus=(req.topic_track == 'core'))}\n"

        "TIME ALLOCATION (use these exact times in section headers; sums to total duration):\n"
        f"{time_table}\n\n"

        "FOLLOW THIS EXACT 9-SECTION STRUCTURE. Each section header MUST be a ## heading with the time in parentheses.\n\n"

        "---\n\n"

        f"## 1. Objective & Competency Anchor ({times[0]} min)\n"
        "- 2-3 SMART learning objectives (Students will be able to ...).\n"
        "- 1-2 NCERT/NCF competencies or LO codes this lesson maps to.\n\n"

        f"## 2. Prior Knowledge Activation ({times[1]} min)\n"
        "- Quick recall task or 2-3 diagnostic questions (with expected answers) probing prerequisites.\n"
        "- If the topic has no real prerequisites, flag the section per the rule above.\n\n"

        f"## 3. Hook & Real-World Connect ({times[2]} min)\n"
        "- A specific, engaging opener (puzzle, story, demo, real-world scenario) — NOT 'discuss the topic'.\n"
        "- 1 essential question and 2-3 student-facing questions with expected answers.\n\n"

        f"## 4. Direct Instruction ({times[3]} min)\n"
        "- Teach each key concept in order. After every concept include a WORKED EXAMPLE (Question → Step-by-Step Solution → Answer).\n"
        "- Include 3-4 worked examples total. Show what to write on the board.\n\n"

        f"## 5. Guided Practice ({times[4]} min)\n"
        "- Named, structured activity (e.g. 'Equation Relay Race'). Steps with sub-step timings.\n"
        "- Group size, student roles, and 4 questions with answers.\n\n"

        f"## 6. Independent Practice ({times[5]} min)\n"
        "- Mixed-difficulty problem set (2 easy, 2 medium, 2 challenging) with answers.\n\n"

        f"## 7. HOTS Task ({times[6]} min)\n"
        "- A higher-order-thinking task: Analyse / Evaluate / Create. Must require justification, not just recall.\n"
        "- Include a short rubric (3 criteria) for what 'good' looks like.\n"
        "- If genuinely inappropriate for the grade/topic, flag the section per the rule above.\n\n"

        f"## 8. Formative Assessment ({times[7]} min)\n"
        "- Exit-ticket-style: 3 specific questions with expected answers.\n"
        "- Note which objective each question evaluates and what mastery looks like.\n\n"

        f"## 9. Closure & Home Connect ({times[8]} min)\n"
        "- 2-sentence student-led recap question, and a single home-connect task that links the lesson to family/community/daily life.\n"
        "- Plus 3-4 homework problems (easy → hard) with answers.\n\n"

        "After the 9 sections, add ONE small **Teacher Notes** section (## Teacher Notes) with: common mistakes (2-3), "
        "time adjustments if running short/long, what the next lesson should cover.\n\n"

        "IMPORTANT: Replace every [bracket] with ACTUAL content. Write real questions with real numbers and real solutions."
    )

    # Use higher token limit for detailed, example-rich output
    lesson_result = call_openai(system_prompt, user_prompt, max_tokens=6000, temperature=0.4)

    # Generate topic overview only if requested (as a separate quick-reference page)
    if req.include_topic_overview:
        topic_system = (
            "You are a subject matter expert. Create a concise teacher quick-reference sheet. "
            f"{lang_profile} "
            "Format with clean markdown: ## for section headers, **bold** for key terms, - for bullet lists."
        )

        topic_prompt = (
            f"Create a 1-page QUICK REFERENCE SHEET for a teacher about to teach this topic:\n\n"
            f"TOPIC: {req.topic}\nSUBJECT: {req.subject}\nGRADE LEVEL: {req.grade_level}\n\n"
            "Include these sections (keep each section concise and useful):\n\n"
            "WHAT THIS TOPIC IS: 2-3 sentence summary of the topic and why students learn it\n\n"
            "KEY FORMULAS / RULES: List all important formulas, theorems, or rules with notation\n\n"
            "ESSENTIAL VOCABULARY: 6-8 key terms with one-line definitions\n\n"
            "COMMON STUDENT MISTAKES: 3-4 specific mistakes students make and how to correct them\n\n"
            "REAL-WORLD CONNECTIONS: 3 specific examples of where this topic appears in real life\n\n"
            "PREREQUISITE CHECK: What students should already know before this lesson\n\n"
            "Use ## for section headers, **bold** for terms, - for bullet points."
        )

        topic_overview = call_openai(topic_system, topic_prompt, max_tokens=2400, temperature=0.3)
        result = (
            "=== TEACHER QUICK REFERENCE ===\n\n"
            + topic_overview
            + "\n\n" + "=" * 50 + "\n\n"
            + "=== LESSON PLAN ===\n\n"
            + lesson_result
        )
    else:
        result = lesson_result

    return {"result": result, "tool": "lesson-plan"}


@app.post("/api/lesson-plan/enrich")
def enrich_lesson_plan(req: LessonPlanEnrichRequest):
    if not req.existing_plan.strip():
        raise HTTPException(status_code=400, detail="Existing plan is required.")
    if not req.topic.strip():
        raise HTTPException(status_code=400, detail="Topic is required.")

    action_map = {
        "more_activities": {
            "title": "ADDITIONAL CLASSROOM ACTIVITIES",
            "instruction": (
                "Generate 4-5 NEW, distinct in-class activities for this lesson that are NOT "
                "already in the existing plan. For each activity include: a clear name, time "
                "estimate, materials needed, step-by-step teacher instructions, student grouping, "
                "and the learning outcome. Mix kinesthetic, collaborative, discussion-based, and "
                "creative formats. Keep them NCERT/CBSE classroom appropriate."
            ),
        },
        "more_examples": {
            "title": "ADDITIONAL NCERT-STYLE EXAMPLE QUESTIONS WITH SOLUTIONS",
            "instruction": (
                "Generate 6-8 ADDITIONAL example questions in NCERT textbook style for the topics "
                "explained in the existing lesson plan. Do NOT repeat any question already shown "
                "above. Cover a range: 2 easy, 3 medium, 2 HOTS (Higher Order Thinking Skills). "
                "For each: state the question, show every solution step with intermediate "
                "reasoning, and end with the final answer. Use the exact tone, notation, and "
                "structure of solved examples in NCERT books for this grade."
            ),
        },
        "more_topics": {
            "title": "ADDITIONAL SUBTOPICS & RELATED TOPICS TO COVER",
            "instruction": (
                "Suggest 5-7 ADDITIONAL subtopics and closely related topics the teacher can "
                "cover to deepen understanding of the main topic. For each subtopic include: "
                "a 1-line description, 2-3 key concepts inside it, suggested time to cover, "
                "1 fully solved NCERT-style example question, and a quick practice question with "
                "answer. Order from foundational to advanced."
            ),
        },
    }

    spec = action_map.get(req.action)
    if not spec:
        raise HTTPException(status_code=400, detail="Unknown action.")

    chapter_context = (
        f"\nNCERT CHAPTER REFERENCE: \"{req.topic_description.strip()}\"\n"
        if req.topic_description.strip() else ""
    )

    system_prompt = (
        "You are a senior NCERT-aligned teacher and curriculum designer. "
        "Extend an existing lesson plan with new, non-overlapping, classroom-ready content.\n\n"
        "OUTPUT FORMAT (STRICT):\n"
        "- Respond in CLEAN MARKDOWN ONLY. No JSON, no LaTeX delimiters ($$ or \\[...]), no code fences around the whole output, no placeholder brackets.\n"
        "- Use ## for section headers, ### for sub-sections, **bold** for key terms, numbered lists for questions, - for bullets.\n"
        "- For math, write equations inline as plain text (e.g. x = (-b ± √(b²-4ac)) / 2a).\n"
        "- Model all questions and solutions on NCERT textbook style with real numbers and worked steps."
    )

    user_prompt = (
        f"Existing lesson plan (for reference — do not repeat its content):\n"
        f"---\n{req.existing_plan[:6000]}\n---\n\n"
        f"TOPIC: {req.topic}\n"
        f"SUBJECT: {req.subject}\n"
        f"GRADE LEVEL: {req.grade_level}\n"
        f"DURATION: {req.duration}\n"
        f"{chapter_context}\n"
        f"Produce a single section titled:\n{spec['title']}\n\n"
        f"{spec['instruction']}\n\n"
        "Output ONLY this new section with clean markdown formatting."
    )

    # "more_topics" requires deeply nested content per subtopic — bump budget.
    enrich_tokens = 3000 if req.action == "more_topics" else 2200
    addition = call_openai(system_prompt, user_prompt, max_tokens=enrich_tokens)
    return {"result": addition, "tool": "lesson-plan", "action": req.action}


@app.post("/api/mc-assessment")
def generate_mc_assessment(req: MCAssessmentRequest):
    if not req.topic.strip():
        raise HTTPException(status_code=400, detail="Topic is required.")

    blooms_map = {
        "remember":   "recall and recognition (define, identify, name, list) — Bloom's Level 1",
        "understand": "comprehension (explain, describe, summarize, classify) — Bloom's Level 2",
        "apply":      "application (use, solve, demonstrate, calculate) — Bloom's Level 3",
        "analyze":    "analysis (compare, contrast, examine, differentiate) — Bloom's Level 4",
        "evaluate":   "evaluation (assess, critique, justify, recommend) — Bloom's Level 5",
        "create":     "synthesis/creation (design, formulate, propose, construct) — Bloom's Level 6",
        "mixed":      "a deliberate mix across all Bloom's Taxonomy levels (recall through synthesis)",
    }

    mc_count = max(1, int(req.num_questions * 0.7))
    other_count = req.num_questions - mc_count

    format_map = {
        "pure_mc":      f"all {req.num_questions} questions as standard multiple choice (4 options: A, B, C, D)",
        "mc_truefalse": f"{mc_count} multiple choice questions (4 options A-D) followed by {other_count} True/False questions, clearly labeled by section",
        "mc_short":     f"{mc_count} multiple choice questions (4 options A-D) followed by {other_count} short answer questions requiring 1-2 sentence responses",
    }

    diff_desc = {
        "easy":   "straightforward recall and basic comprehension — accessible to most students",
        "medium": "moderate application and analysis — requires solid understanding of core concepts",
        "hard":   "challenging evaluation and synthesis — requires deep mastery and critical thinking",
        "mixed":  "a range from basic recall to challenging analysis across difficulty levels",
    }

    answer_key_note = (
        "In the ANSWER KEY, for each question provide: the correct answer letter AND a 1-2 sentence explanation of "
        "why it is correct and what common misconception the wrong options address."
        if req.include_explanations else
        "In the ANSWER KEY, list only the question number and correct answer letter."
    )

    mc_material_note = ""
    if req.source_material.strip():
        mc_material_note = (
            f"\n\nTEACHER-UPLOADED SOURCE MATERIAL (base ALL questions on this content):\n"
            f"---\n{req.source_material[:5000]}\n---\n"
        )

    lang_profile = get_grade_language_profile(req.grade_level)

    system_prompt = (
        "You are an expert CBSE assessment designer with 15+ years of experience writing board-quality papers. "
        "You create rigorous, fair assessments with SPECIFIC, SCENARIO-BASED questions — never generic recall.\n\n"
        "DISTRACTOR QUALITY (CRITICAL):\n"
        "- Every wrong option MUST reflect a specific, realistic student misconception or computational error "
        "(e.g. forgot to convert units, used wrong formula, made sign error, confused two similar concepts).\n"
        "- Wrong options must be the SAME LENGTH and SAME GRAMMATICAL FORM as the correct one — never give the answer away by formatting.\n"
        "- Never use 'all of the above', 'none of the above', double negatives, or trivially wrong options.\n"
        "- If the curriculum-specific detail is unclear, prioritise logical and real-world accuracy over inventing exam formats.\n\n"
        "OUTPUT FORMAT (STRICT):\n"
        "- Respond in CLEAN MARKDOWN ONLY. No JSON, no LaTeX delimiters, no leftover [brackets].\n"
        "- Use # for title, ## for sections, **bold** for key terms, numbered lists for questions, - for bullets.\n"
        f"{lang_profile} "
        + ("ALL questions must come directly from the provided source material. " if req.source_material.strip() else "")
    )

    user_prompt = (
        f"Create a {req.num_questions}-question assessment on: '{req.topic}'\n\n"
        f"GRADE LEVEL: {req.grade_level}\n"
        f"{'SUBJECT: ' + req.subject if req.subject else ''}\n"
        f"DIFFICULTY: {diff_desc.get(req.difficulty, diff_desc['medium'])}\n"
        f"BLOOM'S TAXONOMY: {blooms_map.get(req.blooms_level, blooms_map['mixed'])}\n"
        f"FORMAT: {format_map.get(req.question_format, format_map['pure_mc'])}\n"
        f"{'STANDARDS: ' + req.standards if req.standards else ''}\n"
        f"{'ADDITIONAL INSTRUCTIONS: ' + req.additional_instructions if req.additional_instructions else ''}\n"
        f"{mc_material_note}\n"
        f"{get_curriculum_guardrails(req.grade_level, req.subject, req.topic)}\n"
        "FORMAT:\n"
        "- Title in ALL CAPS with topic and grade\n"
        "- Name / Date / Score fields\n"
        "- Brief student instructions (time allowed, marking scheme)\n"
        "- Questions numbered consecutively\n"
        "- MC options labeled A, B, C, D\n\n"
        "QUESTION QUALITY RULES:\n"
        "- Use real-world scenarios and applied problems (e.g., 'A ladder leans against a wall "
        "at 60 degrees...' not 'What is the sine of an angle?')\n"
        "- Include diagram descriptions where helpful (e.g., 'In the figure, triangle ABC has...')\n"
        "- Mix question types: some conceptual, some calculation-based, some application\n"
        "- Each wrong option should represent a COMMON STUDENT MISTAKE\n\n"
        f"ANSWER KEY: {answer_key_note}\n"
        "For calculation questions, show the working in the answer key.\n\n"
        "Use markdown: # for title, ## for sections, **bold** for important terms, numbered lists for questions."
    )

    # Scale token budget with question count — MCQ + 4 distractors + explanation ~ 350-450 tokens.
    qcount = max(1, int(req.num_questions or 10))
    mc_tokens = min(8000, max(4500, 700 + qcount * 380))
    result = call_openai(system_prompt, user_prompt, max_tokens=mc_tokens)
    return {"result": result, "tool": "mc-assessment"}


# ─── CLASS ACTIVITY GENERATOR ────────────────────────

class ClassActivityRequest(BaseModel):
    topic: str
    grade_level: str
    subject: str = ""
    activity_type: str = "group"
    activity_types: list[str] = []
    num_activities: int = 3
    duration: str = "30 minutes"
    group_size: str = "4-5 students"
    blooms_level: str = "understand"
    learning_outcomes: str = ""
    materials_available: str = ""
    additional_instructions: str = ""
    source_material: str = ""
    topic_description: str = ""
    topic_track: str = "core"

@app.post("/api/class-activity")
def generate_class_activity(req: ClassActivityRequest):
    if not req.topic.strip():
        raise HTTPException(status_code=400, detail="Topic is required.")

    activity_map = {
        "group":       "collaborative group activities where students work together in teams",
        "project":     "project-based exercises that produce a tangible deliverable or presentation",
        "hands_on":    "hands-on kinesthetic activities using physical materials or experiments",
        "discussion":  "structured discussion activities like Socratic seminars, debates, or think-pair-share",
        "game":        "educational games and gamified learning activities",
        "creative":    "creative expression activities like art, drama, writing, or multimedia projects",
    }

    blooms_map = {
        "remember":   "Remember — recall facts, terms, basic concepts",
        "understand": "Understand — explain ideas or concepts in own words",
        "apply":      "Apply — use information in new situations / solve problems",
        "analyze":    "Analyze — draw connections, compare, contrast, break down ideas",
        "evaluate":   "Evaluate — justify a decision, judge information, critique",
        "create":     "Create — produce new or original work, design, compose",
    }

    # Multi-select activity types take priority over the legacy single value.
    selected_types = [t for t in (req.activity_types or []) if t in activity_map]
    if not selected_types and req.activity_type in activity_map:
        selected_types = [req.activity_type]
    if not selected_types:
        selected_types = ["group"]
    types_description = "; ".join(activity_map[t] for t in selected_types)
    blooms_descriptor = blooms_map.get((req.blooms_level or "understand").lower(), blooms_map["understand"])

    scope_block = ""
    if req.topic_description.strip():
        scope_label = "OFFICIAL CBSE/NCERT CHAPTER SCOPE" if req.topic_track == "core" else "TOPIC SCOPE"
        scope_block = (
            f"\n{scope_label}: \"{req.topic_description.strip()}\"\n"
            "ALL activities MUST stay strictly within this scope. Do NOT introduce concepts outside it.\n"
        )

    lang = get_grade_language_profile(req.grade_level)

    material_note = ""
    if req.source_material.strip():
        material_note = (
            f"\n\nTEACHER-UPLOADED SOURCE MATERIAL (align activities to this content):\n"
            f"---\n{req.source_material[:5000]}\n---\n"
        )

    # Subject-specific activity intelligence
    subj_lower = (req.subject or '').lower()
    subject_activity_hints = ""
    if 'math' in subj_lower:
        subject_activity_hints = (
            "MATH ACTIVITY INTELLIGENCE:\n"
            "- Include hands-on manipulatives (paper folding for geometry, blocks for fractions, etc.)\n"
            "- Design activities where students DISCOVER mathematical concepts through exploration\n"
            "- Include real measurement tasks (measure classroom objects, calculate areas)\n"
            "- Use math puzzles, pattern recognition, and logical reasoning games\n"
            "- For higher grades: include real-world modeling (budgeting, statistics from surveys)\n"
        )
    elif any(s in subj_lower for s in ['science', 'physics', 'chemistry', 'biology']):
        subject_activity_hints = (
            "SCIENCE ACTIVITY INTELLIGENCE:\n"
            "- Design actual experiments students can perform with available materials\n"
            "- Include hypothesis → procedure → observation → conclusion format\n"
            "- Use inquiry-based learning where students discover concepts\n"
            "- Include field observation activities (plant growth, weather tracking)\n"
            "- For chemistry/physics: include safe demonstrations with household items\n"
        )
    elif 'english' in subj_lower:
        subject_activity_hints = (
            "ENGLISH ACTIVITY INTELLIGENCE:\n"
            "- Include role-play, dramatization, and reader's theater\n"
            "- Design collaborative writing exercises (story relay, group essay)\n"
            "- Include debate and structured discussion formats\n"
            "- Use creative activities (poetry slam, comic strip creation, news broadcast)\n"
        )
    elif any(s in subj_lower for s in ['history', 'social', 'geography', 'civics']):
        subject_activity_hints = (
            "SOCIAL SCIENCE ACTIVITY INTELLIGENCE:\n"
            "- Include map-making and geographical modeling activities\n"
            "- Design timeline creation and historical reenactment activities\n"
            "- Use case studies from current events connected to topics\n"
            "- Include mock parliament, mock court, or community planning exercises\n"
        )
    elif any(s in subj_lower for s in ['artificial intelligence', 'computer', 'information technology', 'coding', 'programming']) or subj_lower.strip() in ('ai', 'it', 'cs'):
        subject_activity_hints = (
            "AI / COMPUTER SCIENCE / IT ACTIVITY INTELLIGENCE:\n"
            "- Design unplugged activities that teach AI/CS concepts WITHOUT computers when needed\n"
            "- Include hands-on activities: building decision trees on paper, playing the 'AI guessing game'\n"
            "- Design 'AI in real life' scavenger hunts (find AI in daily life — Google Maps, Netflix, Siri)\n"
            "- Include ethical debate activities: 'Should AI replace teachers?', 'Is facial recognition fair?'\n"
            "- Design data collection and analysis activities (survey classmates, create charts, find patterns)\n"
            "- Include role-play: students act as AI components (input, processing, output, feedback loop)\n"
            "- For coding topics: design pair programming, code review, or debugging challenge activities\n"
            "- Include design thinking: 'Design an AI solution for your school' type activities\n"
            "- NEVER design activities that are just math problems with AI/tech words substituted in\n"
            "- Activities should help students understand HOW AI works, not just memorize definitions\n"
        )
    elif 'hindi' in subj_lower:
        subject_activity_hints = (
            "HINDI ACTIVITY INTELLIGENCE:\n"
            "- Include kavita-paath (poetry recitation) and natak (drama) activities\n"
            "- Design collaborative story writing (kahani relay) activities\n"
            "- Include debate (vaad-vivaad) and structured discussion formats\n"
            "- Use creative activities: patra lekhan, poster making, radio natak\n"
        )

    system_prompt = (
        "You are a master teacher and CBSE-trained instructional designer with expertise in "
        "active learning, NEP 2020 pedagogy, and experiential education. "
        "You create activities that are SPECIFIC, ACTIONABLE, and directly tied to curriculum outcomes. "
        "Every activity you design has been tested in real Indian classrooms and works with commonly available materials. "
        "You never write vague activities — every step is detailed, timed, and has clear student roles.\n\n"
        "DEFAULT MATERIALS POLICY (NON-NEGOTIABLE):\n"
        "- Use ONLY locally available, low-cost, everyday materials by default — paper, pencils, chalk, chart paper, "
        "newspapers, used cardboard, bottle caps, pebbles, leaves, string, rubber bands, classroom furniture, and "
        "the human body itself (claps, movement, voice). Items every Indian government / aided school already has.\n"
        "- NEVER suggest premium materials (laminated cards, branded kits, store-bought science kits, expensive "
        "tech) unless the teacher explicitly lists them under MATERIALS AVAILABLE.\n"
        "- This rule applies even if the teacher's ADDITIONAL INSTRUCTIONS is empty — it is your default behaviour, "
        "not a request to be asked for.\n\n"
        "DURATION POLICY (NON-NEGOTIABLE):\n"
        "- The DURATION PER ACTIVITY value the teacher provided is a HARD constraint. Sum of timed sub-steps in "
        "each activity MUST equal that duration (±10%). Do NOT design activities that overflow the time slot.\n\n"
        "SCOPE POLICY:\n"
        "- If a CBSE/NCERT chapter scope or topic scope is given, every activity must stay strictly within that "
        "scope. Do not pull in concepts from later chapters or unrelated topics.\n\n"
        "BLOOM'S ALIGNMENT:\n"
        "- Each activity's cognitive demand must align with the Bloom's level given. Verbs in learning outcomes "
        "and tasks must match (e.g. 'design / create' for Create; 'recall / list' for Remember).\n\n"
        f"{lang} "
        + ("When source material is provided, align ALL activities directly to that content. " if req.source_material.strip() else "")
        + "Format output with CLEAN MARKDOWN for beautiful readability:\n"
        "- Use # for activity number (e.g., # Activity 1: Shape Scavenger Hunt)\n"
        "- Use ## for section headers (e.g., ## Learning Outcomes, ## Materials Needed, ## Procedure)\n"
        "- Use ### for sub-sections\n"
        "- Use **bold** for important labels, terms, and student roles\n"
        "- Use numbered lists (1. 2. 3.) for steps and procedures\n"
        "- Use bullet points (- ) for materials, outcomes, and reflection questions\n"
        "- Use --- between activities as dividers\n"
    )

    user_prompt = (
        f"Create {req.num_activities} detailed, classroom-tested activities for {req.grade_level} students.\n\n"
        f"TOPIC: {req.topic}\n"
        f"{'SUBJECT: ' + req.subject if req.subject else ''}\n"
        f"ACTIVITY TYPES (mix freely across these — at least one activity per selected type when {req.num_activities} >= number of types): {types_description}\n"
        f"DURATION PER ACTIVITY (HARD constraint): {req.duration}\n"
        f"GROUP SIZE: {req.group_size}\n"
        f"BLOOM'S LEVEL TARGET: {blooms_descriptor}\n"
        f"{scope_block}"
        f"{'LEARNING OUTCOMES TO ADDRESS: ' + req.learning_outcomes if req.learning_outcomes else ''}\n"
        f"{'MATERIALS AVAILABLE: ' + req.materials_available if req.materials_available else '(none specified — default to locally available, low-cost everyday materials only)'}\n"
        f"{'ADDITIONAL INSTRUCTIONS: ' + req.additional_instructions if req.additional_instructions else ''}\n"
        f"{material_note}\n\n"
        f"{get_curriculum_guardrails(req.grade_level, req.subject, req.topic, req.topic_description or '', strict_syllabus=(req.topic_track == 'core'))}\n"
        f"{subject_activity_hints}\n"
        "ACTIVITY QUALITY RULES (CRITICAL):\n"
        "- Every activity must be SPECIFIC to the topic — not a generic 'discuss in groups' activity\n"
        "- Include exact materials with quantities (e.g., '4 chart papers, 12 chalk pieces, 1 measuring tape per group')\n"
        "- Materials MUST be locally available and low-cost unless the teacher explicitly listed otherwise\n"
        "- Time every sub-step. Sum of sub-steps in each activity = duration above (±10%)\n"
        "- Define clear student roles within groups (recorder, presenter, material manager, etc.)\n"
        "- Include a specific assessment rubric or checklist for each activity\n"
        "- Align task verbs and assessment with the Bloom's level above\n\n"
        "FOR EACH ACTIVITY, INCLUDE ALL OF THE FOLLOWING:\n\n"
        "1. ACTIVITY TITLE - Creative, engaging name that hints at the concept\n"
        "2. LEARNING OUTCOMES - 2-3 specific SWBAT outcomes aligned to NCERT/CBSE\n"
        "3. MATERIALS NEEDED - Complete list with exact quantities per group\n"
        "4. TEACHER PREPARATION - What to prepare BEFORE class (5-10 min)\n"
        "5. ACTIVITY PROCEDURE - Minute-by-minute breakdown:\n"
        "   - Hook/Introduction (2-3 min) — specific opening question or demonstration\n"
        "   - Main activity steps (numbered, timed)\n"
        "   - Student roles within groups\n"
        "   - Expected student output/deliverable\n"
        "6. DIFFERENTIATION\n"
        "   - Scaffolding: specific support for struggling learners\n"
        "   - Extension: challenge task for advanced students\n"
        "7. ASSESSMENT - Specific rubric or checklist (not just 'observe students')\n"
        "8. REFLECTION QUESTIONS - 2-3 thought-provoking debrief questions\n\n"
        "Make each activity UNIQUE and progressively more challenging. "
        "Activities must be practical for Indian schools with standard resources.\n\n"
        "FORMATTING (CRITICAL):\n"
        "Use clean markdown throughout:\n"
        "- # Activity 1: [Title] for each activity heading\n"
        "- ## for section headers within each activity\n"
        "- **bold** for key labels and student roles\n"
        "- Numbered lists for procedure steps\n"
        "- Bullet points (- ) for materials, outcomes, questions\n"
        "- --- between activities"
    )

    # Each activity needs ~ 750-900 tokens (procedure + materials + diff + rubric + reflection).
    # Underbidding silently truncates the assessment + reflection sections of later activities.
    n = max(1, int(req.num_activities or 3))
    max_tok = min(9000, max(3500, 800 + n * 850))

    result = call_openai(system_prompt, user_prompt, max_tokens=max_tok)
    return {"result": result, "tool": "class-activity"}


# ─── RAG ENDPOINTS (ADVANCED FEATURES) ─────────────────

@app.post("/api/advanced/explain")
async def explain_advanced(request: CodeExplainRequest):
    """Explain advanced code with RAG-retrieved knowledge base"""
    if not RAG_AVAILABLE:
        raise HTTPException(status_code=503, detail="RAG system not available")

    try:
        agent = RAGAgent()
        explanation = agent.explain_code(request.code, request.language, request.grade)
        return {"explanation": explanation, "tool": "explain-code"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/advanced/debug")
async def debug_code(request: DebugRequest):
    """Debug code using pattern matching from knowledge base"""
    if not RAG_AVAILABLE:
        raise HTTPException(status_code=503, detail="RAG system not available")

    try:
        agent = RAGAgent()
        solution = agent.debug_code(request.code, request.language, request.error)
        return {"solution": solution, "tool": "debug-code"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/advanced/improve")
async def improve_code(request: ImproveRequest):
    """Suggest code improvements based on best practices"""
    if not RAG_AVAILABLE:
        raise HTTPException(status_code=503, detail="RAG system not available")

    try:
        agent = RAGAgent()
        improvements = agent.suggest_improvements(request.code, request.language, request.focus)
        return {"improvements": improvements, "tool": "improve-code"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/advanced/analyze")
async def analyze_code(request: AnalyzeRequest):
    """Analyze code structure, complexity, and best practice violations"""
    if not RAG_AVAILABLE:
        raise HTTPException(status_code=503, detail="RAG system not available")

    try:
        analyzer = CodeAnalyzer()
        if request.language == "python":
            analysis = analyzer.analyze_python(request.code)
        elif request.language == "javascript":
            analysis = analyzer.analyze_javascript(request.code)
        else:
            raise HTTPException(status_code=400, detail=f"Language {request.language} not supported")

        return {
            "analysis": analysis,
            "tool": "analyze-code"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/advanced/teach-pattern")
async def teach_pattern(request: PatternRequest):
    """Teach design pattern with code examples"""
    if not RAG_AVAILABLE:
        raise HTTPException(status_code=503, detail="RAG system not available")

    try:
        agent = RAGAgent()
        explanation = agent.teach_pattern(request.pattern, request.language, request.grade)
        return {"explanation": explanation, "tool": "teach-pattern"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ─── CHAT HISTORY & USAGE ENDPOINTS ───────────────────

print("[STARTUP] Loading chat history & usage endpoints...")

class FullHistoryRequest(BaseModel):
    teacher_id: str
    limit: int = 50
    offset: int = 0
    tool_filter: str = ""

@app.post("/api/chat-history")
async def get_chat_history_endpoint(request: ChatHistoryRequest):
    """Get last 7 chats for a teacher (legacy)"""
    try:
        result = get_chat_history(request.teacher_id)
        return result
    except Exception as e:
        print(f"Error getting chat history: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/history")
async def get_full_history(request: FullHistoryRequest):
    """Get full paginated chat history for a teacher"""
    try:
        chats = db.get_chat_history(
            request.teacher_id,
            limit=request.limit,
            offset=request.offset,
            tool_filter=request.tool_filter or None
        )
        return {"success": True, "chats": chats, "count": len(chats)}
    except Exception as e:
        print(f"Error getting history: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/check-usage")
async def check_usage_endpoint(request: UsageCheckRequest):
    """Check daily usage limit for a tool"""
    try:
        result = db.get_usage(request.teacher_id, request.tool_name)
        return result
    except Exception as e:
        print(f"❌ Error checking usage: {e}")
        import traceback
        traceback.print_exc()
        return {
            "usage_count": 0,
            "limit": 50,
            "remaining": 50,
            "exceeded": False
        }

@app.post("/api/increment-usage")
async def increment_usage_endpoint(request: UsageIncrementRequest):
    """Increment usage count and check if limit exceeded"""
    try:
        result = increment_usage(request.teacher_id, request.tool_name)

        # If limit exceeded, return 429 (Too Many Requests) but still return the data
        if result['exceeded']:
            return result  # Frontend handles the error display

        return result
    except Exception as e:
        print(f"❌ Error incrementing usage: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/save-chat")
async def save_chat_endpoint(request: SaveChatRequest):
    """Save a chat to history"""
    try:
        chat_id = db.save_chat(
            request.teacher_id,
            request.tool_name,
            request.topic,
            request.grade_level,
            request.subject,
            request.request_data,
            request.response_preview,
            request.response_content
        )

        return {
            "success": True,
            "chat_id": chat_id,
            "message": "Chat saved to history"
        }
    except Exception as e:
        print(f"❌ Error saving chat: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ─── ADAPTIVE LEARNING ENDPOINTS ───────────────────────

class AdaptiveAssessmentRequest(BaseModel):
    student_id: str
    question_id: int
    teacher_id: str
    answer: str
    is_correct: bool
    time_taken: float
    difficulty_rating: float

@app.post("/api/adaptive/submit-answer")
async def submit_answer(request: AdaptiveAssessmentRequest):
    """Record student answer and update adaptive learning models"""
    try:
        assessment_id = db.record_assessment(
            student_id=request.student_id,
            question_id=request.question_id,
            teacher_id=request.teacher_id,
            answer=request.answer,
            is_correct=request.is_correct,
            time_taken=request.time_taken,
            difficulty_rating=request.difficulty_rating
        )

        return {
            "success": True,
            "assessment_id": assessment_id
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/adaptive/student-progress")
async def get_student_progress(data: dict):
    """Get student's current learning progress"""
    try:
        student_id = data.get("student_id")
        if not student_id:
            return {"success": False, "error": "student_id required"}

        progress = db.get_student_progress(student_id)
        return {
            "success": True,
            "student_id": student_id,
            "progress": progress
        }
    except Exception as e:
        default_progress = {
            "overall_mastery": 0.0,
            "objectives": [],
            "total_topics": 0
        }
        return {
            "success": True,
            "student_id": data.get("student_id", "unknown"),
            "progress": default_progress
        }

@app.post("/api/adaptive/recommend-next")
async def recommend_next(data: dict):
    """Get adaptive learning recommendations for student"""
    try:
        from ml_models import path_recommender

        student_id = data.get("student_id")
        num_recommendations = data.get("num_recommendations", 3)

        if not student_id:
            raise HTTPException(status_code=400, detail="student_id required")

        recommendations = path_recommender.get_recommendations(
            student_id=student_id,
            num_recommendations=num_recommendations
        )

        return {
            "success": True,
            "student_id": student_id,
            "recommendations": recommendations
        }
    except Exception as e:
        default_recommendations = [
            {"topic": "Introduction to the Topic", "reason": "foundational", "priority": 0.9, "difficulty": "easy"},
            {"topic": "Key Concepts and Vocabulary", "reason": "foundational", "priority": 0.8, "difficulty": "easy"},
            {"topic": "Practical Applications", "reason": "foundational", "priority": 0.7, "difficulty": "medium"},
        ]
        return {
            "success": True,
            "student_id": data.get("student_id", "unknown"),
            "recommendations": default_recommendations,
            "note": "Using default recommendations"
        }

@app.post("/api/adaptive/generate-adaptive-question")
async def generate_adaptive_question(data: dict):
    """Generate question at appropriate difficulty level"""
    try:
        from ml_models import difficulty_adaptor, irt_model

        student_id = data.get("student_id")
        topic = data.get("topic")
        grade_level = data.get("grade_level")

        if not all([student_id, topic, grade_level]):
            raise HTTPException(status_code=400, detail="student_id, topic, grade_level required")

        # Get student progress to estimate ability
        progress = db.get_student_progress(student_id)
        student_ability = irt_model.estimate_ability(
            sum(obj['correct_answers'] for obj in progress['objectives']),
            sum(obj['total_attempts'] for obj in progress['objectives'])
        ) if progress['objectives'] else 0.5

        # Suggest appropriate difficulty
        suggested_difficulty = difficulty_adaptor.get_next_difficulty(
            student_ability=student_ability,
            current_difficulty=0.5,
            recent_performance=[]
        )

        # Generate question using existing AI endpoint
        prompt = f"""Generate a {topic} question for {grade_level} level.
        Difficulty: {suggested_difficulty} (0=easy, 1=hard)
        Include 4 multiple choice options and mark the correct answer."""

        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert educator. Generate a clear, engaging educational question."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=300
        )

        question_text = response.choices[0].message.content

        return {
            "success": True,
            "student_id": student_id,
            "topic": topic,
            "difficulty": suggested_difficulty,
            "question": question_text,
            "student_ability": student_ability
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/adaptive/teacher-insights")
async def get_teacher_insights(data: dict):
    """Get analytics and insights for teacher dashboard"""
    try:
        class_id = data.get("class_id")

        if not class_id:
            raise HTTPException(status_code=400, detail="class_id required")

        # Get all students for this class from database
        import sqlite3
        from pathlib import Path
        db_path = Path(__file__).parent / "classroom.db"
        conn = sqlite3.connect(str(db_path))
        c = conn.cursor()

        c.execute('''
            SELECT student_id, name FROM students WHERE teacher_id = ? LIMIT 100
        ''', (class_id,))

        students = c.fetchall()

        if not students:
            return {
                "success": True,
                "stats": {
                    "total_students": 0,
                    "average_mastery": 0.0,
                    "students_below_threshold": 0,
                    "students_above_80": 0,
                    "new_students_today": 0
                },
                "students": [],
                "topics": [],
                "recommendations": []
            }

        # Aggregate student progress data
        student_progresses = []
        topic_mastery_map = {}
        total_mastery = 0
        below_threshold = 0
        above_80 = 0

        for student_id, name in students:
            c.execute('''
                SELECT language, mastery_level, attempts_made, correct_answers FROM learning_objectives
                WHERE student_id = ?
            ''', (student_id,))

            objectives = c.fetchall()

            if objectives:
                student_mastery = sum(obj[1] for obj in objectives) / len(objectives)
                total_attempts = sum(obj[2] for obj in objectives)

                student_progresses.append({
                    'name': name,
                    'student_id': student_id,
                    'overall_mastery': student_mastery,
                    'total_attempts': total_attempts
                })

                total_mastery += student_mastery

                if student_mastery < 0.7:
                    below_threshold += 1
                if student_mastery >= 0.8:
                    above_80 += 1

                # Track topic mastery
                for lang, mastery, attempts, correct in objectives:
                    if lang not in topic_mastery_map:
                        topic_mastery_map[lang] = {'total': 0, 'count': 0}
                    topic_mastery_map[lang]['total'] += mastery
                    topic_mastery_map[lang]['count'] += 1

        avg_mastery = total_mastery / len(student_progresses) if student_progresses else 0

        # Convert topic mastery map to list
        topics = [
            {
                'topic': topic,
                'average_mastery': data['total'] / data['count'] if data['count'] > 0 else 0
            }
            for topic, data in sorted(
                topic_mastery_map.items(),
                key=lambda x: x[1]['total'] / x[1]['count'] if x[1]['count'] > 0 else 0
            )
        ]

        # Get pending recommendations
        c.execute('''
            SELECT student_id, recommended_language, reasoning, difficulty_level, priority_score
            FROM recommendations
            WHERE student_id IN (SELECT student_id FROM students WHERE teacher_id = ?)
            AND status = 'pending'
            ORDER BY priority_score DESC
            LIMIT 10
        ''', (class_id,))

        recommendations = []
        for rec in c.fetchall():
            student_id, lang, reasoning, difficulty, priority = rec
            # Find student name
            student_name = next((s[1] for s in students if s[0] == student_id), f"Student {student_id[:8]}")
            recommendations.append({
                'student_id': student_id,
                'student_name': student_name,
                'recommended_language': lang,
                'reasoning': reasoning,
                'difficulty_level': difficulty,
                'priority': priority
            })

        conn.close()

        return {
            "success": True,
            "stats": {
                "total_students": len(students),
                "average_mastery": avg_mastery,
                "students_below_threshold": below_threshold,
                "students_above_80": above_80,
                "new_students_today": 0
            },
            "students": student_progresses[:20],  # Limit to top 20
            "topics": topics,
            "recommendations": recommendations
        }
    except Exception as e:
        print(f"[ERROR] Teacher insights error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─── INTERACTIVE QUIZ GENERATOR (v2) ──────────────────

class QuizRequest(BaseModel):
    topic: str
    grade_level: str
    subject: str = ""
    num_questions: int = 5
    difficulty: str = "medium"
    question_category: str = "ncert"          # ncert | miscellaneous | advanced
    question_type: str = "mcq"                # mcq | subjective | mix
    paper_mode: bool = False
    topic_description: str = ""
    topic_track: str = "core"                 # core | misc


def _get_subject_question_patterns(subject: str, qtype: str) -> str:
    """Return subject-specific question pattern guidance so the AI generates real exam-quality questions."""
    subj = subject.lower()

    if 'math' in subj:
        if qtype == 'subjective' or qtype == 'mix':
            return (
                "MATHEMATICS SUBJECTIVE QUESTION PATTERNS (MANDATORY — follow these patterns):\n"
                "You MUST generate questions from these categories. Mix them across the set:\n"
                "  1. PROVE / SHOW THAT: Theorem proofs, geometric proofs, algebraic identities\n"
                "     Example pattern: 'Prove that the tangent at any point of a circle is perpendicular to the radius through the point of contact.'\n"
                "  2. FIND / CALCULATE: Numerical problems with given values, requiring step-by-step calculation\n"
                "     Example pattern: 'In a circle with center O, a tangent PT touches the circle at T. If OT = 6 cm and OP = 10 cm, find the length of PT.'\n"
                "  3. CONSTRUCT / DRAW: Geometric constructions with specific measurements\n"
                "     Example pattern: 'Draw a circle of radius 4 cm. From a point 7 cm from its center, construct the pair of tangents to the circle.'\n"
                "  4. APPLICATION / WORD PROBLEM: Real scenario requiring mathematical modeling\n"
                "     Example pattern: 'Two concentric circles are of radii 5 cm and 9 cm. Find the length of the chord of the larger circle which touches the smaller circle.'\n"
                "  5. PROVE WITH CONDITIONS: Problems with given figure/conditions to prove a relationship\n"
                "     Example pattern: 'From an external point P, two tangents PA and PB are drawn to a circle with center O. Prove that PA = PB.'\n\n"
                "CRITICAL RULES FOR MATH:\n"
                "- NEVER generate 'Define X' or 'Explain X' or 'What is X' type questions for math\n"
                "- Every question MUST require mathematical working — calculation, proof, or construction\n"
                "- Include specific numerical values (cm, degrees, etc.) in find/calculate questions\n"
                "- For proofs, state exactly what needs to be proved\n"
                "- For constructions, give exact measurements\n"
                "- At least 40% questions should be PROVE type, 40% FIND/CALCULATE type\n"
            )
        else:  # MCQ
            return (
                "MATHEMATICS MCQ PATTERNS:\n"
                "- Questions must require actual calculation or reasoning, not just recall\n"
                "- Include numerical problems where students must compute the answer\n"
                "- Options should be close numerical values (e.g., 6 cm, 8 cm, 10 cm, 12 cm) to test precision\n"
                "- NEVER ask 'What is the definition of X' as MCQ — ask 'What is the value of X given Y'\n"
            )

    elif 'science' in subj or 'physics' in subj or 'chemistry' in subj or 'biology' in subj:
        if qtype == 'subjective' or qtype == 'mix':
            return (
                "SCIENCE SUBJECTIVE QUESTION PATTERNS (MANDATORY):\n"
                "  1. EXPLAIN MECHANISM: How does X process work? Describe with steps/stages\n"
                "  2. NUMERICAL: Given values, calculate using formulas (for Physics/Chemistry)\n"
                "  3. DIAGRAM-BASED: Draw and label a diagram of X. Explain the function of each part.\n"
                "  4. COMPARE & CONTRAST: Differentiate between X and Y with examples\n"
                "  5. REASON-BASED: Why does X happen? Give scientific reasoning with examples\n"
                "  6. EXPERIMENT: Describe an experiment to demonstrate X. Include materials, procedure, observation, conclusion.\n"
                "  7. APPLICATION: How is X used in daily life / industry? Give 3 examples with explanation.\n\n"
                "CRITICAL: Do NOT generate vague questions like 'Explain photosynthesis' — instead ask:\n"
                "'Describe the light-dependent reactions of photosynthesis. What role does chlorophyll play in capturing solar energy?'\n"
                "- Make questions SPECIFIC with clear scope\n"
                "- Include numerical data where applicable\n"
            )
        else:
            return (
                "SCIENCE MCQ PATTERNS:\n"
                "- Include diagram-based reasoning questions\n"
                "- Numerical problems with calculations\n"
                "- Conceptual questions that test understanding, not memorization\n"
            )

    elif 'english' in subj:
        if qtype == 'subjective' or qtype == 'mix':
            return (
                "ENGLISH SUBJECTIVE QUESTION PATTERNS:\n"
                "  1. PASSAGE-BASED: Provide a short passage and ask comprehension + inference questions\n"
                "  2. CHARACTER ANALYSIS: Analyze a character's motivations, actions, and development\n"
                "  3. CREATIVE WRITING: Write a letter/essay/story on a given topic with word limit\n"
                "  4. GRAMMAR APPLICATION: Rewrite sentences, transform voice/tense, correct errors in context\n"
                "  5. CRITICAL THINKING: Compare themes across texts, evaluate author's perspective\n"
                "- Do NOT ask 'Define noun' or 'What is a verb' — ask questions requiring APPLICATION\n"
            )
        else:
            return "ENGLISH MCQ: Test grammar application, reading comprehension, vocabulary in context.\n"

    elif 'social' in subj or 'history' in subj or 'geography' in subj or 'civics' in subj or 'economics' in subj:
        if qtype == 'subjective' or qtype == 'mix':
            return (
                "SOCIAL SCIENCE SUBJECTIVE PATTERNS:\n"
                "  1. ANALYZE: Why did event X happen? What were its causes and consequences?\n"
                "  2. COMPARE: Compare the policies/movements/features of X and Y\n"
                "  3. MAP-BASED: Locate and explain the significance of places/routes/regions\n"
                "  4. SOURCE-BASED: Read the source and answer — interpret historical documents\n"
                "  5. EVALUATE: Assess the impact of X on society/economy/politics with examples\n"
                "  6. CASE STUDY: Given a scenario, apply concepts to analyze the situation\n"
                "- NEVER ask 'Who was X' or 'When did X happen' — ask 'Analyze WHY X happened and its impact'\n"
            )
        else:
            return "SOCIAL SCIENCE MCQ: Test cause-effect, map skills, source interpretation — not just dates/names.\n"

    elif 'artificial intelligence' in subj or 'ai' == subj.strip().lower() or 'computer' in subj or 'information technology' in subj or 'it' == subj.strip().lower() or 'coding' in subj or 'programming' in subj:
        if qtype == 'subjective' or qtype == 'mix':
            return (
                "AI / COMPUTER SCIENCE / IT SUBJECTIVE QUESTION PATTERNS (MANDATORY):\n"
                "You are generating questions about technology, AI, or computer science. These are NOT math questions.\n"
                "DO NOT generate mathematical computation questions. Generate questions about TECHNOLOGY CONCEPTS.\n\n"
                "  1. REAL-WORLD APPLICATION: How is AI/technology used in a specific real-world domain?\n"
                "     Example: 'Explain how AI is used in healthcare for early disease detection. Describe one specific AI system used in hospitals and how it helps doctors make better diagnoses.'\n"
                "  2. SCENARIO-BASED: Present a real scenario and ask students to apply AI/tech concepts\n"
                "     Example: 'A school wants to build a smart attendance system using AI. Describe what type of AI technology would be needed, what data it would use, and what ethical concerns should be addressed.'\n"
                "  3. ETHICAL ANALYSIS: Analyze ethical implications of AI/technology in society\n"
                "     Example: 'Self-driving cars must make split-second decisions. Discuss the ethical dilemma an AI faces when it must choose between two harmful outcomes. Who should be responsible?'\n"
                "  4. COMPARE TECHNOLOGIES: Compare different AI approaches, tools, or technologies\n"
                "     Example: 'Compare rule-based AI systems with machine learning-based AI systems. Give one real-world example where each approach works better.'\n"
                "  5. DESIGN / CREATE: Design an AI solution for a given problem\n"
                "     Example: 'Design an AI-powered chatbot for a school library. Describe what questions it should answer, what data it needs, and how students would interact with it.'\n"
                "  6. IMPACT ANALYSIS: Evaluate the impact of AI/technology on jobs, education, environment\n"
                "     Example: 'How is AI changing the way farmers grow crops in India? Describe two AI applications in agriculture and their impact on crop yield.'\n\n"
                "CRITICAL RULES FOR AI/CS/IT:\n"
                "- Questions MUST be about technology concepts, applications, ethics, and real-world use cases\n"
                "- DO NOT generate math computation questions disguised as AI questions\n"
                "- DO NOT just replace object names (like 'robots' instead of 'apples') in math problems\n"
                "- Questions should test understanding of HOW technology works, WHERE it is applied, and WHY it matters\n"
                "- Include real company names, real AI tools, real-world scenarios (Google AI, ChatGPT, Tesla Autopilot, etc.)\n"
                "- Focus on: Machine Learning basics, Neural Networks concept, NLP, Computer Vision, Robotics, IoT, Cybersecurity, Data Science, Ethics of AI\n"
                "- Every answer should demonstrate understanding of TECHNOLOGY, not mathematics\n"
            )
        else:  # MCQ
            return (
                "AI / COMPUTER SCIENCE / IT MCQ PATTERNS:\n"
                "- Questions must test understanding of AI/technology concepts, NOT mathematical computation\n"
                "- Include questions about real-world AI applications (healthcare, agriculture, transport, education)\n"
                "- Test knowledge of AI types (supervised/unsupervised learning, NLP, computer vision, robotics)\n"
                "- Include scenario-based questions where students identify the correct AI technique\n"
                "- Ask about ethical considerations, data privacy, bias in AI\n"
                "- Reference real tools and technologies (Python, Scratch, Google AI, ChatGPT, etc.)\n"
                "- NEVER ask mathematical computation questions disguised as AI/tech questions\n"
                "- Options should be technology concepts, not numbers\n"
            )

    elif 'hindi' in subj:
        if qtype == 'subjective' or qtype == 'mix':
            return (
                "HINDI SUBJECTIVE QUESTION PATTERNS:\n"
                "  1. PASSAGE-BASED: Provide a Hindi passage and ask comprehension + inference questions\n"
                "  2. CREATIVE WRITING: Write a letter/essay/story/poem in Hindi on a given topic\n"
                "  3. GRAMMAR APPLICATION: Rewrite sentences, transform voice/tense, sandhi-viched, samas\n"
                "  4. LITERARY ANALYSIS: Analyze themes, characters, and messages in Hindi literature\n"
                "  5. CRITICAL THINKING: Discuss the relevance of a Hindi literary work in modern context\n"
                "- Questions should test language application, not just memory\n"
            )
        else:
            return "HINDI MCQ: Test grammar application, reading comprehension, vocabulary, and literary knowledge.\n"

    # Generic fallback for other subjects
    if qtype == 'subjective' or qtype == 'mix':
        return (
            "SUBJECTIVE QUESTION PATTERNS:\n"
            "IMPORTANT: Generate questions that are GENUINELY about the subject and topic specified.\n"
            "DO NOT generate math questions for non-math subjects. Match the question style to the actual subject.\n\n"
            "  1. ANALYZE / EVALUATE: Questions requiring critical thinking and reasoning about the ACTUAL topic\n"
            "  2. APPLY: Use concepts to solve real-world problems related to the SPECIFIC subject\n"
            "  3. COMPARE: Detailed comparison with examples and reasoning within the subject domain\n"
            "  4. CREATE / DESIGN: Design a solution, experiment, or creative piece relevant to the topic\n"
            "  5. JUSTIFY / PROVE: Support a statement with evidence and logical reasoning from the subject\n"
            "- NEVER generate 'Define X', 'What is X', 'List features of X' type questions\n"
            "- Every question must require THINKING, not just REMEMBERING\n"
            "- Questions must be GENUINELY about the topic — not math problems with topic-related words substituted in\n"
        )
    return ""


def _build_category_prompt(category: str, grade: str, subject: str, topic: str, difficulty: str) -> str:
    """Return category-specific instructions for the AI."""
    if category == "ncert":
        return (
            f"IMPORTANT — NCERT / CBSE BOARD EXAM LEVEL QUESTIONS:\n"
            f"You are generating questions for {grade} {subject} on the topic '{topic}'.\n"
            f"These questions MUST match the style and rigor of actual CBSE board exam papers and NCERT exercise questions.\n\n"
            f"DIFFICULTY CALIBRATION for '{difficulty}':\n"
            f"  - easy: NCERT in-text questions and basic exercise questions. Still require working/reasoning — NOT definitions.\n"
            f"  - medium: NCERT exercise questions and CBSE board exam questions. Application-based, multi-step.\n"
            f"  - hard: HOTS questions from CBSE board papers, exemplar problems, and NCERT exemplar. Require proof, multi-concept integration, or complex calculation.\n\n"
            f"ABSOLUTELY FORBIDDEN:\n"
            f"  - 'Define X', 'What is X', 'Explain X in brief', 'List the properties of X'\n"
            f"  - Any question answerable in one line without mathematical/logical working\n"
            f"  - Generic textbook-style recall questions\n"
            f"  - Questions that just ask to 'compare' or 'differentiate' without specific context\n\n"
            f"EVERY question must require the student to THINK, CALCULATE, PROVE, CONSTRUCT, or ANALYZE.\n"
            f"Model your questions after actual CBSE board papers — look at past year question papers for reference."
        )
    elif category == "miscellaneous":
        return (
            f"MISCELLANEOUS / BROAD KNOWLEDGE QUESTIONS on '{topic}' ({subject}, {grade}):\n"
            f"Generate questions from real-world applications, cross-disciplinary connections, interesting problems, "
            f"and practical scenarios related to this topic.\n"
            f"Questions should go BEYOND the textbook but remain age-appropriate for {grade}.\n"
            f"For '{difficulty}' difficulty:\n"
            f"  - easy: Real-world application problems with given data.\n"
            f"  - medium: Cross-subject connections, practical engineering/daily life scenarios.\n"
            f"  - hard: Research-level thinking, open-ended analysis, multi-domain integration.\n\n"
            f"FORBIDDEN: 'Define X', 'What is X', generic recall questions.\n"
            f"Every question must present a SCENARIO or PROBLEM to solve."
        )
    else:  # advanced
        return (
            f"ADVANCED / COMPETITIVE EXAM LEVEL QUESTIONS on '{topic}' ({subject}, {grade}):\n"
            f"Generate questions at the level of competitive exams appropriate for {grade}:\n"
            f"  - Grade 1-5: Math/Science Olympiad (SOF), IMO, NSO style\n"
            f"  - Grade 6-8: NTSE Stage 1, Regional Olympiad style\n"
            f"  - Grade 9-10: NTSE Stage 2, Pre-RMO, KVPY-SA, Foundation IIT/NEET style\n"
            f"  - Grade 11-12: JEE Main/Advanced, NEET, KVPY, BITSAT style\n"
            f"  - College: GATE, NET, GRE Subject style\n\n"
            f"For '{difficulty}' difficulty:\n"
            f"  - easy: Standard competitive exam questions — one concept, clear approach.\n"
            f"  - medium: Multi-step problems combining 2-3 concepts.\n"
            f"  - hard: Olympiad-level tricky problems, unconventional approaches, proof-based.\n\n"
            f"FORBIDDEN: Textbook-level questions, definitions, simple recall.\n"
            f"Every question must be a genuine PROBLEM that requires deep thinking."
        )


def _build_type_schema(qtype: str, paper_mode: bool) -> str:
    """Return the JSON schema the AI must follow based on question type."""
    mcq_obj = (
        '    {{\n'
        '      "id": 1,\n'
        '      "type": "mcq",\n'
        '      "question": "<Question text>",\n'
        '      "options": ["A) <option>", "B) <option>", "C) <option>", "D) <option>"],\n'
        '      "correct": "A",\n'
        '      "explanation": "<1-2 sentence explanation>",\n'
        '      "marks": 1{section}\n'
        '    }}'
    )
    subj_obj = (
        '    {{\n'
        '      "id": 1,\n'
        '      "type": "subjective",\n'
        '      "question": "<Question text>",\n'
        '      "correct": "",\n'
        '      "options": [],\n'
        '      "explanation": "<Model answer / key points — 2-4 sentences>",\n'
        '      "marks": 2{section}\n'
        '    }}'
    )
    section_field = ',\n      "section": "Section A"' if paper_mode else ''
    mcq_obj = mcq_obj.replace('{section}', section_field)
    subj_obj = subj_obj.replace('{section}', section_field)

    if qtype == "mcq":
        example = mcq_obj
    elif qtype == "subjective":
        example = subj_obj
    else:  # mix
        example = mcq_obj + ',\n' + subj_obj

    paper_fields = ""
    if paper_mode:
        paper_fields = (
            '\n  "paper_mode": true,'
            '\n  "duration": "<suggested time, e.g. 1 Hour>",'
            '\n  "instructions": ["<instruction 1>", "<instruction 2>", "<instruction 3>"],'
        )

    return (
        f'{{\n'
        f'  "title": "<Quiz / Paper title>",{paper_fields}\n'
        f'  "questions": [\n{example}\n  ]\n'
        f'}}'
    )


def _build_type_rules(qtype: str, paper_mode: bool, num_q: int, subject: str = "") -> str:
    """Return rules for the question type."""
    subj_lower = subject.lower()
    is_math = 'math' in subj_lower
    is_tech = any(x in subj_lower for x in ['artificial intelligence', 'computer', 'information technology', 'coding', 'programming']) or subj_lower.strip() in ('ai', 'it', 'cs')

    rules = []
    if qtype == "mcq" or qtype == "mix":
        if is_tech:
            rules += [
                '- For MCQ: "correct" must be A, B, C, or D',
                '- All 4 MCQ options must be plausible technology concepts — NOT numbers or math values',
                '- MCQ marks: 1 mark each',
                '- MCQ questions MUST test understanding of technology concepts, applications, or scenarios',
            ]
        else:
            rules += [
                '- For MCQ: "correct" must be A, B, C, or D',
                '- All 4 MCQ options must be plausible — close values that require actual reasoning to distinguish',
                '- MCQ marks: 1 mark each',
                '- MCQ questions MUST require deep reasoning, NEVER just recall',
            ]
    if qtype == "subjective" or qtype == "mix":
        if is_tech:
            rules += [
                '- For Subjective: "explanation" must contain a COMPLETE model answer (3-6 sentences minimum) about technology concepts',
                '- Marks distribution: 2-mark questions (30%), 3-mark questions (40%), 5-mark questions (30%)',
                '- 2-mark: Short application question about how a specific AI/tech tool works — requires 3-4 lines',
                '- 3-mark: Scenario-based question analyzing real-world AI application or ethical concern — requires 5-8 lines',
                '- 5-mark: Design an AI solution, compare technologies, or analyze societal impact of AI — requires 10+ lines',
                '- ABSOLUTELY NO math computation questions disguised as AI/tech questions',
                '- Every subjective question MUST require the student to demonstrate understanding of TECHNOLOGY concepts',
                '- For model answers in "explanation": explain the technology concept, give real examples, discuss implications',
            ]
        elif is_math:
            rules += [
                '- For Subjective: "explanation" must contain a COMPLETE model answer / solution steps (3-6 sentences minimum)',
                '- Marks distribution: 2-mark questions (30%), 3-mark questions (40%), 5-mark questions (30%)',
                '- 2-mark: Short proof, simple numerical, reason-based — requires 3-4 lines of working',
                '- 3-mark: Application problem, multi-step numerical, prove with diagram — requires 5-8 lines',
                '- 5-mark: Complex proof, construction + proof, long numerical, case-based — requires 10+ lines',
                '- ABSOLUTELY NO "Define X", "What is X", "Explain X", "List features" type questions',
                '- Every subjective question MUST require the student to show mathematical working, write a proof, draw a construction, or solve a multi-step problem',
                '- For model answers in "explanation": show the actual solution steps, not just a summary',
            ]
        else:
            rules += [
                '- For Subjective: "explanation" must contain a COMPLETE model answer (3-6 sentences minimum)',
                '- Marks distribution: 2-mark questions (30%), 3-mark questions (40%), 5-mark questions (30%)',
                '- 2-mark: Short analysis or application question — requires 3-4 lines of thoughtful writing',
                '- 3-mark: Scenario-based or compare/contrast question — requires 5-8 lines',
                '- 5-mark: Detailed analysis, case study, or evaluation — requires 10+ lines',
                '- ABSOLUTELY NO "Define X", "What is X", "List features" type questions',
                f'- Every subjective question MUST be genuinely about the specified subject and topic',
                '- For model answers in "explanation": provide detailed reasoning, examples, and analysis',
            ]
    if qtype == "mix":
        mcq_count = max(1, num_q // 3)
        subj_count = num_q - mcq_count
        rules.append(f'- Generate approximately {mcq_count} MCQ and {subj_count} Subjective questions')
        rules.append('- Put all MCQs first, then subjective questions')

    if paper_mode:
        rules += [
            '- Organize into sections: "Section A – Objective" (1 mark each), "Section B – Short Answer" (2-3 marks), "Section C – Long Answer" (5 marks)',
            '- Include a "section" field on every question',
            '- Provide 3-4 clear exam instructions',
            '- Set appropriate duration based on total marks',
        ]

    rules += [
        '- Every question MUST have "type" set to "mcq" or "subjective"',
        '- QUALITY CHECK: If a question can be answered by just writing a definition or one-line fact, DISCARD it and write a better question',
        '- Include specific numerical values, measurements, and conditions in problems',
        '- Vary the cognitive level: some application, some analysis, some evaluation, some creation',
    ]
    return '\n'.join(rules)


@app.post("/api/quiz")
async def generate_quiz(request: QuizRequest):
    if not request.topic.strip():
        raise HTTPException(status_code=400, detail="Topic cannot be empty")

    grade_profile = get_grade_language_profile(request.grade_level)
    category_prompt = _build_category_prompt(
        request.question_category, request.grade_level,
        request.subject, request.topic, request.difficulty
    )
    subject_patterns = _get_subject_question_patterns(request.subject, request.question_type)
    json_schema = _build_type_schema(request.question_type, request.paper_mode)
    type_rules = _build_type_rules(request.question_type, request.paper_mode, request.num_questions, request.subject)

    topic_ctx = ""
    if request.topic_description:
        topic_ctx = f"\nCurriculum context: {request.topic_description}\n"

    # Detect if this is a math/numerical subject vs conceptual subject
    subj_lower = request.subject.lower()
    is_math = 'math' in subj_lower
    is_science_numerical = any(x in subj_lower for x in ['physics', 'chemistry'])
    is_tech = any(x in subj_lower for x in ['artificial intelligence', 'computer', 'information technology', 'coding', 'programming']) or subj_lower.strip() in ('ai', 'it', 'cs')

    if is_math:
        system_msg = (
            "You are a senior CBSE board exam paper setter with 20+ years of experience in Mathematics. "
            "You write questions EXACTLY like they appear on real printed CBSE math exam papers — complete sentences with variables, measurements, and clear instructions on what to prove/find/construct. "
            "You NEVER write lazy shorthand questions. You NEVER write 'Define X' or 'Explain Y' type questions. "
            "Every question requires genuine mathematical working, logical proof, or step-by-step calculation. "
            "Always respond with valid JSON only. No markdown, no code fences, no extra text outside the JSON."
        )
    elif is_tech:
        system_msg = (
            f"You are a senior exam paper setter specializing in {request.subject} education for school students. "
            f"You have deep knowledge of how AI, machine learning, computer science, and technology work in the REAL WORLD. "
            f"You write questions that test genuine understanding of technology concepts, real-world applications, ethical implications, and practical scenarios. "
            f"You NEVER generate math computation questions disguised as technology questions. "
            f"You NEVER just replace object names (like 'robots' instead of 'apples') in math problems. "
            f"Your questions are about HOW technology works, WHERE it is applied, WHY it matters, and WHAT are its implications. "
            f"Include references to real AI systems (Google AI, ChatGPT, Tesla Autopilot, facial recognition, recommendation systems, etc.). "
            "Always respond with valid JSON only. No markdown, no code fences, no extra text outside the JSON."
        )
    else:
        system_msg = (
            f"You are a senior CBSE board exam paper setter with 20+ years of experience in {request.subject}. "
            f"You write questions EXACTLY like they appear on real printed exam papers — complete, well-formed sentences. "
            f"You NEVER write lazy shorthand questions. Every question requires genuine thinking, analysis, or application. "
            f"Your questions are GENUINELY about {request.subject} — not math problems with subject-related words substituted in. "
            "Always respond with valid JSON only. No markdown, no code fences, no extra text outside the JSON."
        )

    # Build subject-appropriate self-check rules
    if is_math:
        quality_checks = (
            "SELF-CHECK before outputting each question — EVERY question MUST pass ALL checks:\n"
            "  1. Would this EXACT question wording appear on an actual CBSE board exam paper? If no, rewrite it.\n"
            "  2. Does this question require at least 3+ lines of mathematical working to answer? If no, make it harder.\n"
            "  3. Is this just a 'define/explain/list' question? If yes, REPLACE it with a prove/find/construct/analyze question.\n"
            "  4. Is the question text a COMPLETE, well-formed sentence with variable names, measurements, and relationships?\n"
            "  5. For NUMERICAL questions: does it specify exact values (e.g., 'OT = 6 cm, OP = 10 cm')? If no, add specific values.\n"
            "  6. For PROVE questions: does it state EXACTLY what to prove?\n"
            "  7. For CONSTRUCTION questions: does it give exact measurements?\n\n"
            "CRITICAL — QUESTION TEXT QUALITY:\n"
            "- Include variable names (P, Q, O, A, B), measurements (cm, degrees), and relationships in the question text.\n"
            "- Minimum question text length: 15 words for 2-mark, 25 words for 3-mark, 35 words for 5-mark questions."
        )
    elif is_tech:
        quality_checks = (
            f"SELF-CHECK before outputting each question — EVERY question MUST pass ALL checks:\n"
            f"  1. Is this question GENUINELY about {request.subject} / {request.topic}? If it's a math problem with tech words, REWRITE it.\n"
            f"  2. Does this question test understanding of HOW technology works or WHERE it is applied? If no, fix it.\n"
            f"  3. Could a student answer this without knowing anything about {request.topic}? If yes, the question is BAD — rewrite it.\n"
            f"  4. Does the question involve a real-world scenario, ethical dilemma, or practical application? If no, add one.\n"
            f"  5. Are the answer options/explanations about technology concepts (not numbers)? If they're just numbers, REWRITE.\n"
            f"  6. Would a {request.subject} teacher approve this question as relevant to their subject? If no, REPLACE it.\n\n"
            f"CRITICAL — TOPIC RELEVANCE for '{request.topic}':\n"
            f"- Every question must be DIRECTLY about {request.topic} in the context of {request.subject}.\n"
            f"- Use real examples: Google Search AI, Netflix recommendations, Siri/Alexa, self-driving cars, AI in farming, AI in hospitals, facial recognition, spam filters, etc.\n"
            f"- Questions should make students THINK about technology's role in society, not calculate numbers."
        )
    else:
        quality_checks = (
            f"SELF-CHECK before outputting each question — EVERY question MUST pass ALL checks:\n"
            f"  1. Is this question GENUINELY about {request.subject} / {request.topic}? Not a disguised math problem?\n"
            f"  2. Does this question require at least 3+ lines of thoughtful writing to answer? If no, make it harder.\n"
            f"  3. Is this just a 'define/explain/list' question? If yes, REPLACE it with an analyze/evaluate/apply question.\n"
            f"  4. Is the question text a COMPLETE, well-formed sentence ready to print on an exam paper?\n"
            f"  5. Does the question test genuine understanding of {request.topic}? If it's generic, make it specific.\n\n"
            f"CRITICAL — QUESTION TEXT QUALITY:\n"
            f"- Each question MUST be written as a complete, grammatically correct sentence.\n"
            f"- Minimum question text length: 15 words for 2-mark, 25 words for 3-mark, 35 words for 5-mark questions."
        )

    prompt = f"""You are setting a {request.difficulty}-level {request.question_type.upper()} assessment on "{request.topic}" for {request.grade_level} {request.subject}.
{"This is a FULL QUESTION PAPER — include sections, marks allocation, and exam instructions." if request.paper_mode else ""}

{category_prompt}

{subject_patterns}

{grade_profile}
{topic_ctx}

{get_curriculum_guardrails(request.grade_level, request.subject, request.topic, request.topic_description or "", strict_syllabus=(request.question_category == "ncert" and request.topic_track == "core"))}

Generate exactly {request.num_questions} high-quality questions. Each question must be genuinely about "{request.topic}" in the context of {request.subject} — testing real understanding of the subject matter. Every one of the {request.num_questions} questions must be DISTINCT — no repeated or reworded questions, and no two questions sharing the same numbers or scenario.

{quality_checks}

Return ONLY valid JSON (no markdown, no code fences):
{json_schema}

Rules:
{type_rules}"""

    # Scale tokens with question count + type. Subjective answers are ~3x longer than MCQ.
    # Underbidding silently truncates the last 3-5 questions, then the repair regex fakes a "complete" result.
    n = max(1, int(request.num_questions or 10))
    if request.question_type == "subjective":
        max_tok = min(8000, max(2500, 500 + n * 320))
    elif request.question_type == "mix":
        max_tok = min(8000, max(2500, 500 + n * 260))
    else:  # mcq
        max_tok = min(8000, max(2000, 400 + n * 180))
    if request.paper_mode:
        max_tok = max(max_tok, 4500)

    try:
        completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt},
            ],
            model=OPENAI_MODEL, temperature=0.6, max_tokens=max_tok,
        )
        import re as _re
        text = completion.choices[0].message.content.strip()
        text = _re.sub(r'^```[a-z]*\s*', '', text)
        text = _re.sub(r'\s*```$', '', text).strip()
        # Normalize smart quotes that LLMs sometimes emit when teachers paste from Google Docs.
        # These break JSON parsing AND corrupt the unescaped-newline repair regex (attempt #3).
        smart_quote_map = {
            "“": '"', "”": '"',  # curly double quotes
            "‘": "'", "’": "'",  # curly single quotes
            "–": "-", "—": "-",  # en-dash, em-dash
            " ": " ",                  # non-breaking space
        }
        for src, dst in smart_quote_map.items():
            text = text.replace(src, dst)

        # Robust JSON repair: LLMs sometimes produce trailing commas, unescaped chars, or truncated output
        def _repair_json(raw: str) -> dict:
            """Try increasingly aggressive repairs to parse LLM JSON."""
            # Attempt 1: direct parse
            try:
                return json.loads(raw)
            except json.JSONDecodeError:
                pass
            # Attempt 2: fix trailing commas before } or ]
            fixed = _re.sub(r',\s*([\]}])', r'\1', raw)
            try:
                return json.loads(fixed)
            except json.JSONDecodeError:
                pass
            # Attempt 3: fix unescaped newlines inside strings
            fixed2 = _re.sub(r'(?<=": ")(.*?)(?=")', lambda m: m.group(0).replace('\n', '\\n'), fixed)
            try:
                return json.loads(fixed2)
            except json.JSONDecodeError:
                pass
            # Attempt 4: truncated JSON — find last complete question object and close the structure
            last_brace = raw.rfind('}')
            if last_brace > 0:
                truncated = raw[:last_brace + 1]
                # Close any open arrays/objects
                open_brackets = truncated.count('[') - truncated.count(']')
                open_braces = truncated.count('{') - truncated.count('}')
                truncated += ']' * open_brackets + '}' * open_braces
                truncated = _re.sub(r',\s*([\]}])', r'\1', truncated)
                try:
                    return json.loads(truncated)
                except json.JSONDecodeError:
                    pass
            # Attempt 5: extract JSON object from any surrounding text
            match = _re.search(r'\{[\s\S]*\}', raw)
            if match:
                extracted = match.group(0)
                extracted = _re.sub(r',\s*([\]}])', r'\1', extracted)
                open_brackets = extracted.count('[') - extracted.count(']')
                open_braces = extracted.count('{') - extracted.count('}')
                extracted += ']' * open_brackets + '}' * open_braces
                try:
                    return json.loads(extracted)
                except json.JSONDecodeError:
                    pass
            raise json.JSONDecodeError("All repair attempts failed", raw, 0)

        data = _repair_json(text)

        # Drop any duplicate questions, then ensure every question has required fields.
        data["questions"] = _dedupe_questions(data.get("questions", []))
        for i, q in enumerate(data.get("questions", [])):
            q["id"] = i + 1
            q.setdefault("type", "mcq" if q.get("options") else "subjective")
            q.setdefault("marks", 1 if q.get("type") == "mcq" else 2)
            q.setdefault("correct", "")
            q.setdefault("explanation", "")
            q.setdefault("options", [])

        if request.paper_mode:
            data["paper_mode"] = True

        return data
    except json.JSONDecodeError:
        # Retry once with lower temperature for more predictable JSON
        try:
            completion2 = client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_msg + "\nCRITICAL: Your previous response had invalid JSON. Return ONLY a single valid JSON object. No trailing commas, no comments, no truncation."},
                    {"role": "user", "content": prompt},
                ],
                model=OPENAI_MODEL, temperature=0.3, max_tokens=max_tok,
            )
            text2 = completion2.choices[0].message.content.strip()
            text2 = _re.sub(r'^```[a-z]*\s*', '', text2)
            text2 = _re.sub(r'\s*```$', '', text2).strip()
            data = _repair_json(text2)
            data["questions"] = _dedupe_questions(data.get("questions", []))
            for i, q in enumerate(data.get("questions", [])):
                q["id"] = i + 1
                q.setdefault("type", "mcq" if q.get("options") else "subjective")
                q.setdefault("marks", 1 if q.get("type") == "mcq" else 2)
                q.setdefault("correct", "")
                q.setdefault("explanation", "")
                q.setdefault("options", [])
            if request.paper_mode:
                data["paper_mode"] = True
            return data
        except Exception:
            raise HTTPException(status_code=500, detail="AI generated an invalid response. Please try again.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── CODE DEBUGGER ────────────────────────────────────

class CodeDebugRequest(BaseModel):
    code: str
    language: str = "auto-detect"

class RunCodeRequest(BaseModel):
    code: str
    language: str = "auto-detect"

class ExplainSimpleRequest(BaseModel):
    errors: list[str]
    fixes: list[str]
    language: str = "Unknown"

class CSChatMessage(BaseModel):
    role: str
    content: str

class CSChatRequest(BaseModel):
    messages: list[CSChatMessage]

def _strip_json_fences(text: str) -> str:
    """Robustly strip markdown code fences around a JSON payload.
    Handles 1-4 backticks, optional 'json' language hint, leading/trailing whitespace."""
    import re
    s = text.strip()
    # Try to extract content between fenced blocks of 3+ backticks.
    m = re.search(r"`{3,}\s*(?:json|JSON)?\s*\n([\s\S]*?)\n`{3,}", s)
    if m:
        return m.group(1).strip()
    # No proper fences — just strip any leading/trailing backticks and the word 'json'.
    s = re.sub(r"^`+\s*(?:json|JSON)?\s*", "", s)
    s = re.sub(r"\s*`+\s*$", "", s)
    return s.strip()


def _call_openai_json(system_prompt: str, user_prompt: str, max_tokens: int = 4096, temperature: float = 0.1) -> dict:
    """Call the LLM forcing JSON output. Repairs/retries if parsing fails.
    Returns a dict — never raises on parse failure (returns {} instead)."""
    try:
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            temperature=temperature,
            max_tokens=max_tokens,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content.strip()
    except Exception:
        # Some providers reject response_format — fall back to plain mode.
        try:
            resp = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt + "\n\nReturn ONLY a valid JSON object — first character '{' and last '}'."},
                    {"role": "user",   "content": user_prompt},
                ],
                temperature=temperature,
                max_tokens=max_tokens,
            )
            raw = resp.choices[0].message.content.strip()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI error: {e}")

    cleaned = _strip_json_fences(raw)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        # Last-resort: try to slice the first {...} block.
        start, end = cleaned.find("{"), cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                return json.loads(cleaned[start:end + 1])
            except json.JSONDecodeError:
                pass
    return {}


@app.post("/api/debug-code")
def debug_code(req: CodeDebugRequest):
    if not req.code.strip():
        raise HTTPException(status_code=400, detail="Code is required.")

    line_count = req.code.count("\n") + 1
    if len(req.code) > 25000 or line_count > 600:
        raise HTTPException(
            status_code=400,
            detail=f"Code too large ({line_count} lines, {len(req.code)} chars). "
                   "Please paste under 600 lines or 25,000 characters at a time."
        )

    lang_hint = req.language if req.language and req.language != "auto-detect" else ""

    # ── STAGE 1: deep bug analysis (no full code rewrite — keeps response small + lets the model focus) ──
    bug_finder_system = (
        "You are an EXPERT static analyser, code reviewer, and CS tutor. Your ONLY job is to find EVERY bug in the given code. "
        "You must be exhaustive — find ALL of these categories:\n"
        "  1. SYNTAX ERRORS — missing colons, brackets, parentheses, indentation, typos in keywords.\n"
        "  2. LOGIC ERRORS — off-by-one (range(len(x)+1), index 0 vs 1), wrong operator (* instead of +), inverted conditions (% 2 == 1 for even), incorrect return values (0 instead of 1 for factorial base case).\n"
        "  3. RUNTIME ERRORS — division by zero, IndexError (lst[len(lst)/2], items[len(text)-i]), KeyError (dict missing key), TypeError (str + int, int(non-numeric), uppercase() vs upper()).\n"
        "  4. INFINITE LOOPS — missing increment inside while, recursion that never reaches base case (recursive_bug(n) → recursive_bug(n)).\n"
        "  5. API/STDLIB MISUSE — open() in 'r' mode then .write(), json.load(path) instead of json.load(file), .uppercase() instead of .upper(), math.sqrt(negative).\n"
        "  6. TYPE BUGS — comparing str to int, age='30' breaks age>18, str(b) on list.\n"
        "  7. DATA BUGS — initializing max with 0 fails for negative-only lists, duplicates kept while claiming to remove them, count_words returning len+1.\n"
        "  8. RESOURCE BUGS — wrong file mode, dict key not initialized before +=1, opening file without context.\n"
        "  9. ALGORITHMIC BUGS — power(b,e) = b*e, temperature_convert missing +32, withdraw doubles amount, get_middle uses / not //.\n"
        " 10. STYLE / ROBUSTNESS — missing input validation, hard-coded edge cases, unused returns.\n\n"
        "RULES:\n"
        "- Walk through the code function by function, line by line. List EVERY bug you find — do not stop early.\n"
        "- For 200+ lines of buggy code expect 15-40 bugs. NEVER return empty arrays unless the code is truly perfect.\n"
        "- Each 'error' must be specific: include the function name and what is wrong (e.g. 'factorial(): base case returns 0, so factorial(n) is always 0 — should return 1').\n"
        "- Each 'fix' (same index) must say exactly WHAT to change in plain English (e.g. 'Change return 0 to return 1 in the base case').\n"
        "- 'explanation' = one short paragraph (2-3 sentences) summarising the headline issues for a Class 10 student.\n"
        "- Detect the language accurately. Use canonical names: 'Python', 'JavaScript', 'C++', 'Java', etc.\n\n"
        "OUTPUT — return a JSON object with EXACTLY these keys:\n"
        '  { "language": "<name>", "errors_found": ["bug 1", "bug 2", ...], '
        '"fixes_applied": ["fix for bug 1", "fix for bug 2", ...], "explanation": "<paragraph>" }\n'
        "errors_found and fixes_applied MUST have the same length and align 1-to-1."
    )
    bug_finder_user = (
        (f"Language hint: {lang_hint}\n\n" if lang_hint else "") +
        f"Find ALL bugs in this code. Do not skip anything. Be exhaustive.\n\nCODE:\n```\n{req.code}\n```"
    )
    stage1 = _call_openai_json(
        bug_finder_system, bug_finder_user,
        max_tokens=4000, temperature=0.1,
    )

    language = str(stage1.get("language") or lang_hint or "Unknown")
    errors = stage1.get("errors_found") or []
    fixes  = stage1.get("fixes_applied") or []
    if not isinstance(errors, list): errors = [str(errors)]
    if not isinstance(fixes, list): fixes = [str(fixes)]
    errors = [str(e).strip() for e in errors if str(e).strip()]
    fixes  = [str(f).strip() for f in fixes  if str(f).strip()]
    while len(fixes) < len(errors):
        fixes.append("Apply the matching correction described above.")
    fixes = fixes[:len(errors) or len(fixes)]
    explanation = str(stage1.get("explanation") or "").strip()

    # ── STAGE 2: produce the corrected file (only runs if bugs were found) ──
    debugged_code = req.code
    if errors:
        bug_summary = "\n".join(f"- {e}" for e in errors[:40])
        fixer_system = (
            "You are a precise code rewriter. The user will give you BUGGY code and a list of bugs. "
            "Your only job: return the FULL corrected file. Apply EVERY listed fix. Keep all variable names and structure "
            "where possible. Preserve comments. Output proper indentation. Output ONLY a JSON object."
        )
        fixer_user = (
            f"Language: {language}\n\n"
            f"BUGS TO FIX:\n{bug_summary}\n\n"
            f"BUGGY CODE:\n```\n{req.code}\n```\n\n"
            'Return JSON: { "debugged_code": "<full corrected file as a single string, with \\n for newlines>" }'
        )
        # Generous tokens — corrected file may grow slightly with safe-guards.
        fix_tokens = min(8000, max(2000, 600 + line_count * 22))
        stage2 = _call_openai_json(fixer_system, fixer_user, max_tokens=fix_tokens, temperature=0.0)
        candidate = str(stage2.get("debugged_code") or "").strip()
        if candidate:
            debugged_code = candidate

    if not errors:
        explanation = explanation or "Your code looks clean — no bugs were found."

    return {
        "original_code": req.code,
        "language": language,
        "errors_found": errors,
        "fixes_applied": fixes,
        "explanation": explanation,
        "debugged_code": debugged_code,
    }


# ─── RUN CODE (executes Python/JS/TS locally; AI-simulates others) ─────────────

def _detect_lang_from_code(code: str) -> str:
    s = code.strip()
    if not s:
        return "Python"
    if "#include" in s or ("int main(" in s and ("printf(" in s or "cout" in s)):
        return "C++" if "cout" in s or "std::" in s else "C"
    if "public static void main" in s or "System.out.println" in s:
        return "Java"
    if s.startswith("<!DOCTYPE") or "<html" in s.lower():
        return "HTML"
    if "def " in s and ("print(" in s or ":" in s):
        return "Python"
    if "console.log(" in s or ("function " in s and "{" in s):
        return "JavaScript"
    if "package main" in s and "func " in s:
        return "Go"
    return "Python"

@app.post("/api/run-code")
def run_code(req: RunCodeRequest):
    import subprocess, tempfile, shutil, re as _re
    if not req.code.strip():
        raise HTTPException(status_code=400, detail="Code cannot be empty")

    code = req.code
    lang = (req.language or "auto-detect").strip()
    if lang.lower() in ("auto-detect", ""):
        lang = _detect_lang_from_code(code)
    lang_lower = lang.lower()

    # Python: real execution
    if lang_lower == "python":
        try:
            with tempfile.NamedTemporaryFile(suffix=".py", mode="w", delete=False, encoding="utf-8") as f:
                f.write(code); tmp_path = f.name
            try:
                proc = subprocess.run(
                    [sys.executable, tmp_path],
                    capture_output=True, text=True, timeout=10,
                    env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"}
                )
                return {"output": proc.stdout[:5000], "error": proc.stderr[:2000],
                        "exit_code": proc.returncode, "language": "Python", "simulated": False}
            finally:
                try: os.unlink(tmp_path)
                except Exception: pass
        except subprocess.TimeoutExpired:
            return {"output": "", "error": "Execution timed out (10s limit)", "exit_code": -1, "language": "Python", "simulated": False}
        except Exception as e:
            return {"output": "", "error": str(e), "exit_code": -1, "language": "Python", "simulated": False}

    # JS / TS via Node
    if lang_lower in ("javascript", "javascript (react)", "typescript", "typescript (react)") and shutil.which("node"):
        run_str = code
        if "typescript" in lang_lower:
            run_str = _re.sub(r":\s*\w+(\[\])?", "", run_str)
            run_str = _re.sub(r"<\w+>", "", run_str)
        try:
            with tempfile.NamedTemporaryFile(suffix=".js", mode="w", delete=False, encoding="utf-8") as f:
                f.write(run_str); tmp_path = f.name
            try:
                proc = subprocess.run(["node", tmp_path], capture_output=True, text=True, timeout=8)
                return {"output": proc.stdout[:5000], "error": proc.stderr[:2000],
                        "exit_code": proc.returncode, "language": lang, "simulated": False}
            finally:
                try: os.unlink(tmp_path)
                except Exception: pass
        except Exception as e:
            return {"output": "", "error": str(e), "exit_code": -1, "language": lang, "simulated": False}

    # Fallback: AI-simulated execution for compiled / unsupported languages.
    # The model MUST be honest — if the code has compile/runtime bugs, return the
    # error in `error` and exit_code 1. Do NOT fabricate successful output.
    try:
        sys_prompt = (
            f"You are an honest {lang} compiler + runtime. Mentally execute the user's code and report what would actually happen. "
            "RULES:\n"
            "- If the code has ANY compile error, syntax error, type error, or runtime error, report it exactly as the real compiler/runtime would.\n"
            "- NEVER fabricate a successful output if the code is broken. Be ruthlessly honest.\n"
            "- Return STRICT JSON only with these keys: "
            '{"status":"ok"|"error", "output":"<stdout text or empty>", "error":"<error message or empty>"}.\n'
            "- If the program would crash before printing anything, output is empty and status='error'.\n"
            "- If it runs cleanly, status='ok' and error is empty."
        )
        result = _call_openai_json(
            sys_prompt,
            f"Run this {lang} code honestly. Report any compile/runtime errors as the real compiler would:\n\n```\n{code}\n```",
            max_tokens=700, temperature=0,
        )
        status = str(result.get("status", "ok")).lower()
        out = str(result.get("output") or "").strip()
        err = str(result.get("error") or "").strip()
        if status == "error" or err:
            return {"output": out, "error": err or "Code failed to run.", "exit_code": 1,
                    "language": lang, "simulated": True}
        return {"output": f"[Simulated output — {lang}]\n{out}".strip() if out else "(no output)",
                "error": "", "exit_code": 0, "language": lang, "simulated": True}
    except Exception as e:
        return {"output": "", "error": f"Could not run {lang}: {e}", "exit_code": -1, "language": lang, "simulated": False}


# ─── EXPLAIN SIMPLY ────────────────────────────────────────────────────────────

@app.post("/api/explain-simple")
def explain_simple(req: ExplainSimpleRequest):
    if not req.errors:
        raise HTTPException(status_code=400, detail="No errors to explain")
    sys_prompt = (
        "You are a friendly coding mentor for school kids. Explain bugs and fixes "
        "in the simplest possible way, like talking to a 10-year-old. Use short sentences, "
        "fun analogies, and be encouraging. Use a few emojis for warmth (not too many)."
    )
    user_prompt = (
        f"A student just had their {req.language} code debugged. "
        "Explain each bug and how it was fixed in the simplest words possible.\n\n"
        "Bugs found:\n" + "\n".join(f"- {e}" for e in req.errors) + "\n\n"
        "Fixes applied:\n" + "\n".join(f"- {f}" for f in req.fixes) + "\n\n"
        "Keep the whole explanation under 6 short sentences total."
    )
    text = call_openai(sys_prompt, user_prompt, max_tokens=500, temperature=0.7)
    return {"explanation": text.strip()}


# ─── CS TUTOR CHAT (CS topics only) ────────────────────────────────────────────

@app.post("/api/cs-tutor")
def cs_tutor(req: CSChatRequest):
    if not req.messages:
        raise HTTPException(status_code=400, detail="No messages provided")
    system_msg = {
        "role": "system",
        "content": (
            "You are CodeVidhya's CS Tutor — a friendly, expert teacher for Computer Science and Programming. "
            "You ONLY answer questions about computer science, coding, programming languages, algorithms, "
            "data structures, software engineering, databases, networks, operating systems, AI/ML basics, "
            "web development, and closely related technical topics. "
            "If asked anything UNRELATED to CS, politely reply: \"I'm CodeVidhya's CS Tutor, so I can only help "
            "with computer science and programming topics. Ask me anything about coding!\" "
            "Give clear, beginner-friendly answers. Use small examples when helpful. "
            "Format with short paragraphs and line breaks for readability. "
            "Never mention which AI model or company powers you. You are CodeVidhya's CS Tutor — that's it."
        )
    }
    msgs = [system_msg] + [{"role": m.role, "content": m.content} for m in req.messages]
    try:
        resp = client.chat.completions.create(
            model=OPENAI_MODEL, messages=msgs, temperature=0.5, max_tokens=900,
        )
        return {"reply": resp.choices[0].message.content.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Tutor error: {e}")


# ─── FEEDBACK GENERATOR ────────────────────────────────

class FeedbackRequest(BaseModel):
    student_name: str
    grade_level: str
    feedback_type: str = "general"
    tone: str = "encouraging"
    ratings: dict | None = None
    context: str = ""

@app.post("/api/generate-feedback")
def generate_feedback(req: FeedbackRequest):
    if not req.student_name.strip():
        raise HTTPException(status_code=400, detail="Student name is required.")
    if not req.grade_level.strip():
        raise HTTPException(status_code=400, detail="Grade level is required.")

    # Validate + summarize ratings so the AI can reason about strengths vs. growth areas.
    rating_lines = ""
    strengths, growth_areas = [], []
    if req.ratings:
        clean = {}
        for k, v in req.ratings.items():
            try:
                iv = int(v)
                if 1 <= iv <= 5:
                    clean[str(k)] = iv
            except (TypeError, ValueError):
                # Silently skip non-numeric values rather than confusing the AI.
                continue
        if clean:
            rating_lines = "\n".join([f"- {k}: {v}/5" for k, v in clean.items()])
            strengths    = [k for k, v in clean.items() if v >= 4]
            growth_areas = [k for k, v in clean.items() if v <= 3]
    if not rating_lines:
        rating_lines = "Not provided"

    tone_guides = {
        "encouraging":  "warm, motivating, future-focused. Celebrate effort and progress.",
        "constructive": "balanced and direct. Specific praise paired with specific growth steps.",
        "formal":       "polished and professional, suitable for a printed report card.",
        "warm":         "personal and caring, like a teacher who knows the student well.",
        "professional": "concise, measured, and observation-based with concrete examples.",
    }
    tone_guide = tone_guides.get((req.tone or "").lower(), tone_guides["encouraging"])

    strengths_hint = f"High-rated areas (write strengths around these first): {', '.join(strengths)}." if strengths else ""
    growth_hint    = f"Lower-rated areas (suggest growth here): {', '.join(growth_areas)}." if growth_areas else ""

    system_prompt = (
        "You are a thoughtful school teacher writing personalised end-of-term feedback for ONE student. "
        "Your feedback feels human, specific, and never generic.\n\n"
        "WRITING RULES:\n"
        "- Address the student by FIRST NAME only.\n"
        "- Length: 110-170 words. One flowing paragraph (not bullet points, not multiple paragraphs).\n"
        "- Structure: (1) Open with a specific strength tied to a behaviour, (2) acknowledge effort or growth area, "
        "(3) one concrete suggestion for next term, (4) warm closing.\n"
        "- NEVER use clichés like 'good job', 'keep it up', 'continue the good work', 'I am happy to say'.\n"
        "- Reference the rated areas naturally — do NOT list ratings or scores in the text.\n"
        "- Match the requested tone exactly.\n\n"
        "OUTPUT FORMAT (STRICT):\n"
        "Return ONLY the feedback paragraph as plain text. No JSON, no bullet points, no markdown headers, no '**bold**' tags, "
        "no preamble like 'Here is the feedback:'. Just the paragraph itself."
    )

    user_prompt = (
        f"Student first name: {req.student_name.split()[0]}\n"
        f"Full name (for context only, do not repeat in feedback): {req.student_name}\n"
        f"Grade level: {req.grade_level}\n"
        f"Feedback focus area: {req.feedback_type}\n"
        f"Tone: {req.tone} — {tone_guide}\n"
        f"Ratings (out of 5, for your reasoning only — do NOT list these in the paragraph):\n{rating_lines}\n"
        f"{strengths_hint}\n{growth_hint}\n"
        f"Extra teacher context: {req.context.strip() or 'None'}\n\n"
        "Write the feedback paragraph now."
    )

    text = call_openai(system_prompt, user_prompt, max_tokens=600, temperature=0.75)
    # Light cleanup: strip stray quotes, markdown, or leading preamble.
    cleaned = text.strip().strip('"').strip("'")
    for prefix in ("Feedback:", "Here is the feedback:", "Here's the feedback:"):
        if cleaned.lower().startswith(prefix.lower()):
            cleaned = cleaned[len(prefix):].strip()
    return {
        "student_name": req.student_name,
        "grade_level": req.grade_level,
        "feedback_type": req.feedback_type,
        "tone": req.tone,
        "generated_feedback": cleaned,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# ─── VOCABULARY MASTERY TOOL ENDPOINTS ──────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

# ── Pydantic models ──

class VocabWorksheetRequest(BaseModel):
    topic: str
    grade_level: int
    learning_objective: str
    source_text: str | None = None
    additional_context: str | None = None
    session_id: str | None = None

class VocabSessionCreate(BaseModel):
    metadata: dict | None = None

class VocabRAGDocRequest(BaseModel):
    content: str
    topic: str | None = ""
    grade_level: int | None = 0


# ── Vocabulary helper functions ──

def _vocab_count_syllables(word: str) -> int:
    word = (word or "").lower().strip(".,!?;:'\"")
    if not word:
        return 1
    count = len(_re_mod.findall(r'[aeiouy]+', word))
    if word.endswith('e') and count > 1:
        count -= 1
    return max(count, 1)

VOCAB_GRADE_VOCAB_CAPS = {
    1: {"max_syllables": 2, "max_chars": 7},
    2: {"max_syllables": 3, "max_chars": 8},
    3: {"max_syllables": 3, "max_chars": 9},
}

def _vocab_check_grade_complexity(data: dict, grade_level: int) -> "str | None":
    caps = VOCAB_GRADE_VOCAB_CAPS.get(grade_level)
    if not caps:
        return None
    too_complex = []
    for vw in data.get("vocab_words", []):
        word = (vw.get("word") or "").strip().lower()
        if not word:
            continue
        if _vocab_count_syllables(word) > caps["max_syllables"] or len(word) > caps["max_chars"]:
            too_complex.append(word)
    if too_complex:
        return (
            f"Grade {grade_level} vocabulary must be at most {caps['max_syllables']} syllables and "
            f"{caps['max_chars']} letters per word. These words are too complex: "
            f"{', '.join(too_complex[:5])}. REPLACE each one with a simpler Grade {grade_level} word "
            f"on the same topic (1-2 syllable sight word, CVC pattern, decodable phonics)."
        )
    return None

def _vocab_validate(data: dict, min_items: int = 8) -> "str | None":
    vocab_words = data.get("vocab_words")
    if not isinstance(vocab_words, list) or len(vocab_words) < min_items:
        return f"vocab_words must have at least {min_items} items, got {len(vocab_words) if isinstance(vocab_words, list) else 'missing'}"
    matching = data.get("matching_section")
    if not isinstance(matching, dict) or not matching.get("items"):
        return "matching_section is missing or has no items"
    fib = data.get("fill_in_blank")
    if not isinstance(fib, dict):
        return "fill_in_blank section is missing"
    sentences = fib.get("sentences")
    if not isinstance(sentences, list) or len(sentences) < min_items:
        return f"fill_in_blank.sentences must have at least {min_items} items, got {len(sentences) if isinstance(sentences, list) else 'missing'}"
    sw = data.get("sentence_writing")
    if not isinstance(sw, dict) or not sw.get("prompts"):
        return "sentence_writing section is missing or has no prompts"
    prompts = sw.get("prompts")
    if not isinstance(prompts, list) or len(prompts) < min_items:
        return f"sentence_writing.prompts must have at least {min_items} items, got {len(prompts) if isinstance(prompts, list) else 'missing'}"
    return None


# Groq fallback model list (shared by both tools)
_TOOL_FALLBACK_MODELS = [
    OPENAI_MODEL,
    "llama-3.1-8b-instant",
    "deepseek-r1-distill-llama-70b",
    "meta-llama/llama-4-scout-17b-16e-instruct",
]


# ── Vocabulary Sessions ──

@app.post("/api/vocabulary/sessions")
async def vocab_new_session(req: VocabSessionCreate):
    session_id = vocab_create_session(req.metadata)
    return {"session_id": session_id}

@app.get("/api/vocabulary/sessions/{session_id}/history")
async def vocab_session_history(session_id: str):
    return {"session_id": session_id, "history": vocab_get_session_history(session_id)}

@app.get("/api/vocabulary/worksheets")
async def vocab_list_worksheets(limit: int = 20):
    return {"worksheets": vocab_get_all_worksheets(limit)}


# ── Vocabulary Generate (SSE streaming) ──

@app.post("/api/vocabulary/generate")
async def vocab_generate_worksheet(req: VocabWorksheetRequest, request: Request):
    from fastapi.responses import StreamingResponse as _SR

    await generate_limiter.check(client_ip(request))
    session_id = req.session_id or vocab_create_session()

    vocab_rag_retriever.build_index()
    rag_context = vocab_rag_retriever.build_context(
        f"{req.topic} grade {req.grade_level} {req.learning_objective}",
        grade_level=req.grade_level,
    )
    grade_ctx = vocab_get_grade_prompt_context(req.grade_level)

    source_block = (
        "\nSOURCE MATERIAL (MANDATORY -- pick vocabulary words from THIS content; "
        "definitions, examples and context sentences must reflect what the source actually says):\n"
        f"---\n{req.source_text[:6000]}\n---\n"
    ) if req.source_text else ""
    additional_block = f"Additional Context: {req.additional_context}" if req.additional_context else ""
    rag_block = f"\n{rag_context}" if rag_context else ""
    ctx_block = f"{source_block}{additional_block}\n{rag_block}".strip()

    def _build_prompt(extra_instructions: str = "") -> str:
        p = VOCAB_GRADE_PROFILES.get(req.grade_level, VOCAB_GRADE_PROFILES[7])
        n = vocab_get_word_count(req.grade_level)

        return f"""You are an expert educator and curriculum specialist.
Your task is to create a grade-calibrated Vocabulary Mastery Worksheet.

{grade_ctx}

CONTENT DETAILS:
Topic: {req.topic}
Learning Objective: {req.learning_objective}
{ctx_block}

CRITICAL RULES:
1. Generate EXACTLY {n} vocabulary words -- this count is calibrated for Grade {req.grade_level} attention span. Do NOT generate more or fewer.
2. ALL {n} vocabulary words must be exactly right for Grade {req.grade_level} students (age {req.grade_level + 5}-{req.grade_level + 6}).
3. Every definition must use SIMPLER words than the target word -- a Grade {req.grade_level} student must understand it.
4. Fill-in-blank sentences must be at Grade {req.grade_level} reading level: {p['sentence']}
5. Sentence writing hints must be Grade {req.grade_level} appropriate: {p['hint_style']}
6. Do NOT use words from other grade levels. Do NOT use placeholder words like 'word1'.
{"7. SOURCE MATERIAL is provided above and is the AUTHORITATIVE basis for this worksheet. EVERY vocabulary word MUST be picked from words that actually appear in the SOURCE MATERIAL -- do NOT introduce vocabulary that is not present in the source. Definitions must reflect how the word is used in the source. Fill-in-blank sentences must mirror the source's topic/context." if req.source_text else ""}
{extra_instructions}

Return ONLY valid JSON. No markdown fences. No prose outside the JSON.

{{
  "vocab_words": [
    {{"word": "actual Grade {req.grade_level} word from topic", "definition": "Grade {req.grade_level}-appropriate definition", "part_of_speech": "noun|verb|adjective|adverb"}},
    ... EXACTLY {n} words total, all relevant to '{req.topic}', all appropriate for Grade {req.grade_level}
  ],
  "matching_section": {{
    "title": "Section 1: Match the Word to Its Meaning",
    "instructions": "Grade {req.grade_level}-appropriate matching instruction (1 sentence).",
    "items": [
      {{"word": "word from vocab_words", "definition": "Grade {req.grade_level}-appropriate definition"}},
      ... all {n} words, definitions in SHUFFLED order (not matching vocab_words order)
    ]
  }},
  "fill_in_blank": {{
    "title": "Section 2: Fill in the Blank",
    "instructions": "Grade {req.grade_level}-appropriate instruction (1 sentence).",
    "word_bank": ["all {n} vocab words listed here"],
    "sentences": [
      {{"sentence": "Grade {req.grade_level} sentence with ___ blank for the answer word.", "answer": "the correct vocab word"}},
      ... EXACTLY {n} sentences total, one for EACH vocabulary word
    ]
  }},
  "sentence_writing": {{
    "title": "Section 3: Write Your Own Sentences",
    "instructions": "Grade {req.grade_level}-appropriate writing instruction.",
    "prompts": [
      {{"word": "vocab word", "hint": "{p['hint_style']}", "example": "Grade {req.grade_level}-appropriate example sentence"}},
      ... EXACTLY {n} prompts total, one for EACH vocabulary word
    ]
  }}
}}"""

    def _sse(obj: dict) -> str:
        return f"data: {json.dumps(obj)}\n\n"

    def stream_gen():
        max_attempts = 5
        extra_instructions = ""
        last_reason = ""
        model_idx = 0

        for attempt in range(1, max_attempts + 1):
            if attempt > 1:
                yield _sse({"type": "retry", "attempt": attempt, "reason": last_reason})

            current_model = _TOOL_FALLBACK_MODELS[min(model_idx, len(_TOOL_FALLBACK_MODELS) - 1)]
            yield _sse({"type": "progress", "message": f"Attempt {attempt}: calling {current_model}..."})

            prompt = _build_prompt(extra_instructions)
            collected_chunks = []

            try:
                stream = client.chat.completions.create(
                    model=current_model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=4500,
                    stream=True,
                )

                for chunk in stream:
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta.content or ""
                    if delta:
                        collected_chunks.append(delta)
                        yield _sse({"type": "token", "content": delta})

            except Exception as exc:
                last_reason = str(exc)
                if model_idx < len(_TOOL_FALLBACK_MODELS) - 1:
                    model_idx += 1
                    next_model = _TOOL_FALLBACK_MODELS[model_idx]
                    yield _sse({"type": "status", "message": f"Model error -- switching to {next_model}..."})
                    extra_instructions = ""
                else:
                    extra_instructions = f"IMPORTANT: Fix the following error from the previous attempt: {last_reason}\n"
                continue

            raw = "".join(collected_chunks).strip()
            for fence in ("```json", "```"):
                if raw.startswith(fence):
                    raw = raw[len(fence):]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
            raw = _re_mod.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', raw)

            yield _sse({"type": "status", "message": "Parsing JSON response..."})

            first, last = raw.find("{"), raw.rfind("}")
            if first != -1 and last != -1 and last > first:
                raw = raw[first:last + 1]

            data = None
            try:
                data = json.loads(raw)
            except json.JSONDecodeError as exc:
                try:
                    from json_repair import repair_json
                    repaired = repair_json(raw)
                    data = json.loads(repaired)
                except Exception:
                    last_reason = f"Invalid JSON: {exc}"
                    extra_instructions = (
                        "CRITICAL: Your previous response was not valid JSON. "
                        "Return ONLY a raw JSON object -- no markdown fences, no prose. "
                        "Escape every double-quote inside string values as \\\". "
                        "Do not put a comment line starting with '...' inside the JSON.\n"
                    )
                    continue

            yield _sse({"type": "status", "message": "Validating worksheet structure..."})

            validation_error = _vocab_validate(data, min_items=vocab_get_word_count(req.grade_level))
            if validation_error:
                last_reason = f"Validation failed: {validation_error}"
                _n = vocab_get_word_count(req.grade_level)
                extra_instructions = (
                    f"IMPORTANT: Fix this validation error from your previous attempt: {validation_error}. "
                    f"Ensure vocab_words has exactly {_n} items, matching_section has {_n} items, "
                    f"fill_in_blank has exactly {_n} sentences, and sentence_writing has exactly {_n} prompts.\n"
                )
                continue

            complexity_error = _vocab_check_grade_complexity(data, req.grade_level)
            if complexity_error:
                last_reason = f"Grade-complexity failed: {complexity_error}"
                yield _sse({"type": "status", "message": "Words too complex for the grade -- regenerating with simpler vocabulary..."})
                extra_instructions = complexity_error + "\n"
                continue

            yield _sse({"type": "status", "message": "Saving worksheet..."})

            full_content = {
                **data,
                "rag_context_used": bool(rag_context),
            }

            try:
                worksheet_id = vocab_save_worksheet(
                    session_id=session_id,
                    topic=req.topic,
                    grade_level=req.grade_level,
                    learning_objective=req.learning_objective,
                    content=full_content,
                )
                vocab_save_rag_document(
                    content=(
                        f"vocabulary worksheet topic {req.topic} grade {req.grade_level} "
                        f"objective {req.learning_objective} words "
                        + " ".join(w["word"] for w in data.get("vocab_words", []))
                    ),
                    doc_type="worksheet",
                    topic=req.topic,
                    grade_level=req.grade_level,
                )
                vocab_rag_retriever.build_index()
            except Exception as exc:
                yield _sse({"type": "error", "message": f"Database error: {exc}"})
                return

            yield _sse({
                "type": "complete",
                "session_id": session_id,
                "worksheet_id": worksheet_id,
                "worksheet": full_content,
            })
            return

        yield _sse({"type": "error", "message": f"Failed after {max_attempts} attempts. Last error: {last_reason}"})

    return _SR(
        stream_gen(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Vocabulary Export DOCX ──

@app.post("/api/vocabulary/export/docx")
async def vocab_export_docx(payload: dict):
    from fastapi.responses import StreamingResponse as _SR
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ws = payload.get("worksheet", {})
    topic = payload.get("topic", "Vocabulary")
    grade = payload.get("grade_level", "")
    objective = payload.get("learning_objective", "")

    doc = Document()
    title = doc.add_heading("Vocabulary Mastery Worksheet", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Topic: {topic}  |  Grade: {grade}  |  Objective: {objective}")
    doc.add_paragraph("Name: ____________________________   Date: _______________")
    doc.add_paragraph()

    matching = ws.get("matching_section", {})
    if matching:
        doc.add_heading(matching.get("title", "Section 1: Matching"), 1)
        doc.add_paragraph(matching.get("instructions", ""))
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Vocabulary Word"
        hdr[1].text = "Definition"
        for item in matching.get("items", []):
            row = tbl.add_row().cells
            row[0].text = item.get("word", "")
            row[1].text = ""
        doc.add_paragraph()

    fib = ws.get("fill_in_blank", {})
    if fib:
        doc.add_heading(fib.get("title", "Section 2: Fill in the Blank"), 1)
        doc.add_paragraph(fib.get("instructions", ""))
        wb = ", ".join(fib.get("word_bank", []))
        doc.add_paragraph(f"Word Bank: [ {wb} ]")
        doc.add_paragraph()
        for i, s in enumerate(fib.get("sentences", []), 1):
            doc.add_paragraph(f"{i}. {s.get('sentence', '')}")
        doc.add_paragraph()

    sw = ws.get("sentence_writing", {})
    if sw:
        doc.add_heading(sw.get("title", "Section 3: Write Your Own Sentences"), 1)
        doc.add_paragraph(sw.get("instructions", ""))
        doc.add_paragraph()
        for i, p in enumerate(sw.get("prompts", []), 1):
            doc.add_paragraph(f"{i}. Word: {p.get('word', '')}")
            doc.add_paragraph(f"   Hint: {p.get('hint', '')}")
            doc.add_paragraph("   My sentence: _________________________________________________")
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return _SR(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="vocabulary_{topic}.docx"'},
    )


# ── Vocabulary RAG document upload ──

@app.post("/api/vocabulary/rag/add-text")
async def vocab_add_rag_text(req: VocabRAGDocRequest):
    doc_id = vocab_save_rag_document(req.content, "knowledge", req.topic, req.grade_level)
    vocab_rag_retriever.build_index()
    return {"success": True, "doc_id": doc_id}

@app.post("/api/vocabulary/rag/add-file")
async def vocab_add_rag_file(request: Request, file: UploadFile = File(...)):
    from fastapi.responses import StreamingResponse as _SR
    await upload_limiter.check(client_ip(request))
    raw = await read_upload_capped(file)
    content = ""
    if file.filename.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(raw))
        content = " ".join(p.extract_text() or "" for p in reader.pages)
    elif file.filename.endswith(".docx"):
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(raw))
        content = " ".join(p.text for p in doc.paragraphs)
    else:
        content = raw.decode("utf-8", errors="ignore")
    content = (content or "").strip()
    doc_id = vocab_save_rag_document(content[:6000], "file", file.filename, 0)
    vocab_rag_retriever.build_index()
    return {
        "success": True,
        "doc_id": doc_id,
        "chars_indexed": len(content),
        "text": content[:8000],
        "filename": file.filename,
    }


# ── Vocabulary URL extraction ──

@app.post("/api/vocabulary/extract-url")
async def vocab_extract_url(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    url = assert_public_url((req.get("url") or "").strip())
    try:
        import httpx
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=20.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; VocabularyTool/1.0)"}) as cx:
            r = await cx.get(url)
            assert_public_url(str(r.url))
            r.raise_for_status()
            soup = BeautifulSoup(r.text, "html.parser")
            for bad in soup(["script", "style", "noscript", "iframe", "nav", "footer", "header", "form", "aside"]):
                bad.decompose()
            title = (soup.title.string or "").strip() if soup.title else ""
            main = soup.find("main") or soup.find("article") or soup.body or soup
            text = _re_mod.sub(r"\s+\n", "\n", main.get_text("\n", strip=True))
            text = _re_mod.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Could not extract readable text from this page.")
        return {"success": True, "title": title, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"URL fetch failed: {e}")


# ── Vocabulary YouTube extraction ──

async def _vocab_youtube_metadata_fallback(video_id: str, url: str) -> dict | None:
    import httpx
    title = ""
    author = ""
    description = ""
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; VocabularyTool/1.0)"}) as cx:
            r = await cx.get("https://www.youtube.com/oembed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}", "format": "json"})
            if r.status_code == 200:
                j = r.json()
                title = (j.get("title") or "").strip()
                author = (j.get("author_name") or "").strip()
    except Exception:
        pass
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; VocabularyTool/1.0)"}) as cx:
            r = await cx.get("https://noembed.com/embed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}"})
            if r.status_code == 200:
                j = r.json()
                if not title:  title  = (j.get("title") or "").strip()
                if not author: author = (j.get("author_name") or "").strip()
                description = (j.get("description") or "").strip()
    except Exception:
        pass
    try:
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                              "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept-Language": "en-US,en;q=0.9",
                "Cookie": "CONSENT=YES+cb.20210328-17-p0.en+FX+000",
            }) as cx:
            r = await cx.get(f"https://www.youtube.com/watch?v={video_id}")
            if r.status_code == 200:
                soup = BeautifulSoup(r.text, "html.parser")
                if not title:
                    ot = soup.find("meta", attrs={"property": "og:title"})
                    if ot: title = (ot.get("content") or "").strip()
                ot_desc = soup.find("meta", attrs={"property": "og:description"})
                if ot_desc and not description:
                    description = (ot_desc.get("content") or "").strip()
                for s in soup.find_all("script"):
                    txt = s.string or ""
                    if "shortDescription" in txt:
                        mm = _re_mod.search(r'"shortDescription":"((?:\\.|[^"\\])*)"', txt)
                        if mm:
                            full = mm.group(1).encode("utf-8").decode("unicode_escape", errors="ignore")
                            if len(full) > len(description):
                                description = full
                            break
    except Exception:
        pass
    pieces = []
    if title:       pieces.append(f"Video title: {title}")
    if author:      pieces.append(f"Channel: {author}")
    if description: pieces.append(f"\n{description}")
    text = "\n".join(pieces).strip()
    if not text or len(text) < 20:
        return None
    return {"title": title or f"YouTube video {video_id}", "text": text}

@app.post("/api/vocabulary/extract-youtube")
async def vocab_extract_youtube(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    url = (req.get("url") or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="url is required")
    m = _re_mod.search(r"(?:v=|youtu\.be/|/embed/|/shorts/)([A-Za-z0-9_-]{11})", url)
    if not m:
        raise HTTPException(status_code=400, detail="Could not detect a YouTube video id in that URL.")
    video_id = m.group(1)
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        proxy_config = None
        wh_user = os.getenv("WEBSHARE_PROXY_USERNAME")
        wh_pass = os.getenv("WEBSHARE_PROXY_PASSWORD")
        if wh_user and wh_pass:
            try:
                from youtube_transcript_api.proxies import WebshareProxyConfig
                proxy_config = WebshareProxyConfig(proxy_username=wh_user, proxy_password=wh_pass)
            except Exception:
                pass
        if proxy_config is not None:
            transcript = YouTubeTranscriptApi(proxy_config=proxy_config).fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        elif hasattr(YouTubeTranscriptApi, "get_transcript"):
            chunks = YouTubeTranscriptApi.get_transcript(video_id)
            text = " ".join((c.get("text", "") if isinstance(c, dict) else getattr(c, "text", "")) for c in chunks).strip()
        else:
            transcript = YouTubeTranscriptApi().fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Transcript was empty.")
        return {"success": True, "video_id": video_id, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except Exception as e:
        msg = str(e)
        blocked = ("blocking requests" in msg.lower()
                   or ("ip" in msg.lower() and "block" in msg.lower())
                   or "could not retrieve a transcript" in msg.lower())
        if blocked:
            fallback = await _vocab_youtube_metadata_fallback(video_id, url)
            if fallback:
                return {
                    "success": True, "video_id": video_id, "url": url,
                    "text": fallback["text"][:8000], "chars": len(fallback["text"]),
                    "title": fallback["title"],
                    "note": "Transcript was blocked -- using title + description instead.",
                }
        raise HTTPException(
            status_code=502 if blocked else 500,
            detail=(
                "YouTube blocks transcript requests from our server. Please paste the transcript manually."
                if blocked else f"YouTube transcript fetch failed: {msg}"
            ),
        )


# ── Vocabulary Auto-fields ──

@app.post("/api/vocabulary/auto-fields")
async def vocab_auto_fields(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    source_text = (req.get("source_text") or "").strip()
    if not source_text:
        raise HTTPException(status_code=400, detail="source_text is required.")
    grade = req.get("grade_level")
    grade_hint = f"Calibrate the topic + objective for Grade {grade} students. " if grade else ""
    prompt = (
        "You are an expert vocabulary curriculum designer.\n"
        "Read the SOURCE MATERIAL below and propose a clean, classroom-ready\n"
        "  * Topic (a short noun phrase, 4-10 words, naming the main idea/theme of the material)\n"
        "  * Learning Objective (one sentence starting with 'Students will...' "
        "describing what vocabulary skill the student will gain).\n"
        f"{grade_hint}"
        "Return ONLY a JSON object with keys 'topic' and 'learning_objective'. "
        "No markdown, no prose outside the JSON.\n\n"
        "SOURCE MATERIAL:\n---\n"
        f"{source_text[:6000]}\n---\n\n"
        'JSON: {"topic": "...", "learning_objective": "Students will ..."}'
    )
    try:
        completion = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You return strict JSON. No markdown fences."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=400,
            response_format={"type": "json_object"},
        )
        raw = completion.choices[0].message.content.strip()
        data = json.loads(raw)
        topic = (data.get("topic") or "").strip()
        objective = (data.get("learning_objective") or "").strip()
        if not topic and not objective:
            raise HTTPException(status_code=502, detail="Model returned no fields.")
        return {"success": True, "topic": topic, "learning_objective": objective}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"auto-fields failed: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# ─── READING COMPREHENSION TOOL ENDPOINTS ───────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

# ── Pydantic models ──

class ReadingComprehensionRequest(BaseModel):
    topic: str
    grade_level: int
    learning_objective: str
    source_text: str | None = None
    additional_context: str | None = None
    session_id: str | None = None

class ReadingSessionCreate(BaseModel):
    metadata: dict | None = None

class ReadingRAGDocRequest(BaseModel):
    content: str
    topic: str | None = ""
    grade_level: int | None = 0

class CompleteAnswerRequest(BaseModel):
    question: str
    passage_text: str
    grade_level: int
    word_limit: int = 35
    question_type: str | None = "literal"
    answer_hint: str | None = ""


# ── Reading helper functions ──

READING_GRADE_VOCAB_CAPS = {
    1: {"max_syllables": 2, "max_chars": 7},
    2: {"max_syllables": 3, "max_chars": 8},
    3: {"max_syllables": 3, "max_chars": 9},
}

def _reading_count_syllables(word: str) -> int:
    word = (word or "").lower().strip(".,!?;:'\"")
    if not word:
        return 1
    count = len(_re_mod.findall(r'[aeiouy]+', word))
    if word.endswith('e') and count > 1:
        count -= 1
    return max(count, 1)

def _reading_check_grade_complexity(data: dict, grade_level: int) -> "str | None":
    caps = READING_GRADE_VOCAB_CAPS.get(grade_level)
    if not caps:
        return None
    vic = data.get("vocabulary_in_context", {})
    too_complex = []
    for item in vic.get("items", []) or []:
        word = (item.get("word") or "").strip().lower()
        if not word:
            continue
        if _reading_count_syllables(word) > caps["max_syllables"] or len(word) > caps["max_chars"]:
            too_complex.append(word)
    if too_complex:
        return (
            f"Grade {grade_level} vocabulary must be at most {caps['max_syllables']} syllables and "
            f"{caps['max_chars']} letters per word. These words are too complex: "
            f"{', '.join(too_complex[:5])}. REPLACE each with a simpler Grade {grade_level} word "
            f"on the same topic (1-2 syllable sight word, CVC pattern, decodable phonics)."
        )
    return None

def _reading_validate(data: dict, grade_level: int = 7) -> "str | None":
    byr = data.get("before_you_read")
    if not isinstance(byr, dict) or not isinstance(byr.get("questions"), list) or len(byr["questions"]) < 1:
        return "before_you_read.questions is missing or empty"
    passage = data.get("passage")
    if not isinstance(passage, dict) or not passage.get("text"):
        return "passage.text is missing or empty"
    try:
        p = READING_GRADE_PROFILES.get(grade_level, READING_GRADE_PROFILES[7])
        rng = p.get("passage_words", "")
        nums = [int(x) for x in _re_mod.findall(r"\d+", rng)]
        if nums:
            max_words = max(nums)
            actual = len(passage["text"].split())
            ceiling = int(max_words * 1.6)
            if actual > ceiling:
                return (f"passage is {actual} words but Grade {grade_level} must be "
                        f"{rng} words. Rewrite it MUCH shorter -- no more than {max_words} words.")
    except Exception:
        pass
    tdq = data.get("text_dependent_questions")
    if not isinstance(tdq, dict) or not isinstance(tdq.get("questions"), list) or len(tdq["questions"]) < 1:
        return "text_dependent_questions.questions is missing or empty"
    vic = data.get("vocabulary_in_context")
    if not isinstance(vic, dict) or not isinstance(vic.get("items"), list) or len(vic["items"]) < 1:
        return "vocabulary_in_context.items is missing or empty"
    return None


# ── Reading Sessions ──

@app.post("/api/reading/sessions")
async def reading_new_session(req: ReadingSessionCreate):
    session_id = reading_create_session(req.metadata)
    return {"session_id": session_id}

@app.get("/api/reading/sessions/{session_id}/history")
async def reading_session_history(session_id: str):
    return {"session_id": session_id, "history": reading_get_session_history(session_id)}

@app.get("/api/reading/comprehensions")
async def reading_list_comprehensions(limit: int = 20):
    return {"comprehensions": reading_get_all_comprehensions(limit)}


# ── Reading Complete Answer ──

@app.post("/api/reading/complete-answer")
async def reading_complete_answer(req: CompleteAnswerRequest):
    p = READING_GRADE_PROFILES.get(req.grade_level, READING_GRADE_PROFILES[7])
    grade_ctx = reading_get_grade_prompt_context(req.grade_level)

    prompt = f"""You are an expert educator writing a model answer for a Grade {req.grade_level} student.

{grade_ctx}

READING PASSAGE:
{req.passage_text[:2000]}

QUESTION TYPE: {req.question_type}
QUESTION: {req.question}

STRICT WORD LIMIT: {req.word_limit} words maximum. Count every word -- do NOT exceed this limit.

TASK: Write a model answer in EXACTLY {req.word_limit} words or fewer.
RULES:
1. HARD LIMIT: Your answer must be {req.word_limit} words or fewer. This is non-negotiable.
2. Use ONLY Grade {req.grade_level} vocabulary: {p['vocab']}
3. Sentence structure for Grade {req.grade_level}: {p['sentence']}
4. Cognitive level: {p['blooms']}
5. Reference the passage text as evidence but stay within the word limit.
6. Write ONLY the answer -- no "Answer:", no labels, no explanation outside the answer.

Answer ({req.word_limit} words max):"""

    last_error = ""
    for model in _TOOL_FALLBACK_MODELS:
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=min(req.word_limit * 3, 300),
            )
            raw_content = (response.choices[0].message.content if response.choices else None) or ""
            answer = raw_content.strip()
            answer = _re_mod.sub(r'^(answer|model answer|response)\s*:\s*', '', answer, flags=_re_mod.IGNORECASE).strip()
            if not answer:
                last_error = f"empty completion from {model}"
                continue
            words = answer.split()
            if len(words) > req.word_limit:
                answer = " ".join(words[:req.word_limit])
                for punct in ('.', '!', '?'):
                    last_punct = answer.rfind(punct)
                    if last_punct > len(answer) // 2:
                        answer = answer[:last_punct + 1]
                        break
            return {"answer": answer}
        except Exception as exc:
            last_error = str(exc)
            continue

    raise HTTPException(status_code=503, detail=f"Failed after {len(_TOOL_FALLBACK_MODELS)} attempts: {last_error}")


# ── Reading Generate (SSE streaming) ──

@app.post("/api/reading/generate")
async def reading_generate_comprehension(req: ReadingComprehensionRequest, request: Request):
    from fastapi.responses import StreamingResponse as _SR

    await generate_limiter.check(client_ip(request))
    session_id = req.session_id or reading_create_session()

    grade_ctx = reading_get_grade_prompt_context(req.grade_level)

    source_block = (
        "\nSOURCE MATERIAL (MANDATORY -- base the passage's facts, names, examples and "
        "vocabulary on THIS content only; do not invent facts that contradict it):\n"
        f"---\n{req.source_text[:6000]}\n---\n"
    ) if req.source_text else ""
    additional_block = f"Additional Context: {req.additional_context}" if req.additional_context else ""
    rag_block = ""
    ctx_block = f"{source_block}{additional_block}\n{rag_block}".strip()

    def _build_prompt(extra_instructions: str = "") -> str:
        p = READING_GRADE_PROFILES.get(req.grade_level, READING_GRADE_PROFILES[7])
        word_range = p["passage_words"]
        c = get_reading_counts(req.grade_level)

        low_grade_block = f"\n- Passage must be {word_range} words.\n"

        return f"""You are an expert reading specialist and curriculum designer.
Your task is to create a complete, grade-calibrated Reading Comprehension activity.

{grade_ctx}
{low_grade_block}
CONTENT DETAILS:
Topic: {req.topic}
Learning Objective: {req.learning_objective}
{ctx_block}

CRITICAL RULES:
1. Every word you write -- passage, questions, instructions, hints -- must match Grade {req.grade_level} level EXACTLY.
2. The passage MUST be {word_range} words -- count carefully.
3. Generate EXACTLY {c['total_q']} text-dependent questions (calibrated for Grade {req.grade_level} attention) and EXACTLY {c['vocab']} vocabulary items. Do NOT write all literal questions.
4. Vocabulary in Context words must come directly from the passage.
5. Before You Read questions must activate prior knowledge at a Grade {req.grade_level} cognitive level.
{"6. SOURCE MATERIAL is provided above and is the AUTHORITATIVE basis for this passage. The passage MUST be a grade-level rewrite/summary of the SOURCE MATERIAL -- every fact, name, number, date, event, term and example must come directly from it. Do NOT invent content from your own knowledge if it contradicts or is absent from the source. Text-dependent questions must reference the rewritten passage (which reflects the source); Vocabulary in Context words must be picked from words actually present in the source." if req.source_text else ""}
{extra_instructions}

Return ONLY valid JSON. No markdown fences. No prose outside the JSON.

{{
  "before_you_read": {{
    "title": "Before You Read",
    "instructions": "Grade {req.grade_level}-appropriate activation prompt here (1 sentence).",
    "questions": [
      {{"number": 1, "question": "Grade {req.grade_level} prior-knowledge question about {req.topic}", "type": "activation"}},
      {{"number": 2, "question": "Grade {req.grade_level} prediction question about the passage", "type": "prediction"}},
      {{"number": 3, "question": "Grade {req.grade_level} inquiry question the student wonders about {req.topic}", "type": "inquiry"}}
    ]
  }},
  "annotation_guide": {{
    "title": "Annotation Guide",
    "instructions": "Grade {req.grade_level}-appropriate reading strategy instruction (1 sentence).",
    "symbols": [
      {{"symbol": "star", "meaning": "Grade {req.grade_level} explanation of main idea marking"}},
      {{"symbol": "?", "meaning": "Grade {req.grade_level} explanation of confusion marking"}},
      {{"symbol": "!", "meaning": "Grade {req.grade_level} explanation of interesting info marking"}},
      {{"symbol": "->", "meaning": "Grade {req.grade_level} explanation of cause-effect marking"}},
      {{"symbol": "circle", "meaning": "Grade {req.grade_level} explanation of vocabulary marking"}}
    ]
  }},
  "passage": {{
    "title": "Engaging title relevant to {req.topic}",
    "text": "Write the FULL passage here. Must be {word_range} words. Use paragraph breaks (\\n\\n). Every sentence must match Grade {req.grade_level} syntax and vocabulary.",
    "word_count": "actual number of words in the passage you wrote"
  }},
  "text_dependent_questions": {{
    "title": "Text-Dependent Questions",
    "instructions": "Grade {req.grade_level}-appropriate instruction for answering with text evidence.",
    "questions": [
      {{"number": 1, "question": "Question at Grade {req.grade_level} level", "type": "literal", "answer_hint": "Paragraph evidence"}},
      ... Generate EXACTLY {c['total_q']} questions total for Grade {req.grade_level}:
          {c['literal']} literal (type "literal"), {c['inferential']} inferential (type "inferential"){', ' + str(c['higher']) + ' higher-order Analyze/Evaluate (type "critical_thinking")' if c['higher'] else ''}.
          Number them 1..{c['total_q']}. Each needs an answer_hint pointing to passage evidence.
    ]
  }},
  "vocabulary_in_context": {{
    "title": "Vocabulary in Context",
    "instructions": "Grade {req.grade_level}-appropriate vocabulary strategy instruction.",
    "items": [
      {{
        "word": "actual word from the passage appropriate for Grade {req.grade_level}",
        "sentence_from_passage": "Copy the exact sentence from your passage containing this word.",
        "context_clue_type": "definition|example|contrast|inference",
        "activity": "Grade {req.grade_level}-appropriate activity using this word",
        "answer": "Grade {req.grade_level}-appropriate answer"
      }},
      ... EXACTLY {c['vocab']} items total, each word from the passage
    ]
  }}
}}"""

    def _sse(obj: dict) -> str:
        return f"data: {json.dumps(obj)}\n\n"

    def stream_gen():
        max_attempts = 5
        extra_instructions = ""
        last_reason = ""
        model_idx = 0

        for attempt in range(1, max_attempts + 1):
            if attempt > 1:
                yield _sse({"type": "retry", "attempt": attempt, "reason": last_reason})

            current_model = _TOOL_FALLBACK_MODELS[min(model_idx, len(_TOOL_FALLBACK_MODELS) - 1)]
            yield _sse({"type": "progress", "message": f"Attempt {attempt}: calling {current_model}..."})

            prompt = _build_prompt(extra_instructions)
            collected_chunks = []

            try:
                stream = client.chat.completions.create(
                    model=current_model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.75,
                    max_tokens=4500,
                    stream=True,
                )

                for chunk in stream:
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta.content or ""
                    if delta:
                        collected_chunks.append(delta)
                        yield _sse({"type": "token", "content": delta})

            except Exception as exc:
                last_reason = str(exc)
                if model_idx < len(_TOOL_FALLBACK_MODELS) - 1:
                    model_idx += 1
                    next_model = _TOOL_FALLBACK_MODELS[model_idx]
                    yield _sse({"type": "status", "message": f"Model error -- switching to {next_model}..."})
                    extra_instructions = ""
                else:
                    extra_instructions = f"IMPORTANT: Fix the following error from the previous attempt: {last_reason}\n"
                continue

            raw = "".join(collected_chunks).strip()
            for fence in ("```json", "```"):
                if raw.startswith(fence):
                    raw = raw[len(fence):]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
            raw = _re_mod.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', raw)

            yield _sse({"type": "status", "message": "Parsing JSON response..."})

            first, last = raw.find("{"), raw.rfind("}")
            if first != -1 and last != -1 and last > first:
                raw = raw[first:last + 1]

            data = None
            try:
                data = json.loads(raw)
            except json.JSONDecodeError as exc:
                try:
                    from json_repair import repair_json
                    repaired = repair_json(raw)
                    data = json.loads(repaired)
                except Exception:
                    last_reason = f"Invalid JSON: {exc}"
                    extra_instructions = (
                        "CRITICAL: Your previous response was not valid JSON. "
                        "Return ONLY a raw JSON object -- no markdown fences, no prose. "
                        "Escape every double-quote inside string values as \\\".\n"
                    )
                    if model_idx < len(_TOOL_FALLBACK_MODELS) - 1:
                        model_idx += 1
                        yield _sse({"type": "status", "message": f"Switching to {_TOOL_FALLBACK_MODELS[model_idx]}..."})
                    continue

            yield _sse({"type": "status", "message": "Validating comprehension structure..."})

            validation_error = _reading_validate(data, req.grade_level)
            if validation_error:
                last_reason = f"Validation failed: {validation_error}"
                extra_instructions = (
                    f"IMPORTANT: Fix this validation error from your previous attempt: {validation_error}. "
                    "Ensure before_you_read has >=3 questions, passage.text is present, "
                    f"passage is {READING_GRADE_PROFILES.get(req.grade_level, READING_GRADE_PROFILES[7])['passage_words']} words, "
                    "text_dependent_questions has >=6 questions, and vocabulary_in_context has >=5 items.\n"
                )
                if model_idx < len(_TOOL_FALLBACK_MODELS) - 1:
                    model_idx += 1
                    yield _sse({"type": "status", "message": f"Switching to {_TOOL_FALLBACK_MODELS[model_idx]}..."})
                continue

            complexity_error = _reading_check_grade_complexity(data, req.grade_level)
            if complexity_error:
                last_reason = f"Grade-complexity failed: {complexity_error}"
                yield _sse({"type": "status", "message": "Words too complex for the grade -- regenerating with simpler vocabulary..."})
                extra_instructions = complexity_error + "\n"
                continue

            # Annotate passage with readability metrics
            passage_text = data.get("passage", {}).get("text", "")
            readability = reading_analyze_text_grade(passage_text)
            if readability:
                data["passage"]["readability"] = readability
                word_count = readability.get("word_count", 0)
                if word_count:
                    data["passage"]["word_count"] = word_count

            yield _sse({"type": "status", "message": "Saving comprehension..."})

            full_content = {**data, "rag_context_used": False}

            try:
                comp_id = reading_save_comprehension(
                    session_id=session_id,
                    topic=req.topic,
                    grade_level=req.grade_level,
                    learning_objective=req.learning_objective,
                    content=full_content,
                )
                reading_save_rag_document(
                    content=(
                        f"reading comprehension topic {req.topic} grade {req.grade_level} "
                        f"objective {req.learning_objective} passage: {passage_text[:300]}"
                    ),
                    doc_type="comprehension",
                    topic=req.topic,
                    grade_level=req.grade_level,
                )
            except Exception as exc:
                yield _sse({"type": "error", "message": f"Database error: {exc}"})
                return

            yield _sse({
                "type": "complete",
                "session_id": session_id,
                "comprehension_id": comp_id,
                "comprehension": full_content,
            })
            return

        yield _sse({"type": "error", "message": f"Failed after {max_attempts} attempts. Last error: {last_reason}"})

    return _SR(
        stream_gen(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Reading Export DOCX ──

@app.post("/api/reading/export/docx")
async def reading_export_docx(payload: dict):
    from fastapi.responses import StreamingResponse as _SR
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    comp = payload.get("comprehension", {})
    topic = payload.get("topic", "Reading")
    grade = payload.get("grade_level", "")
    objective = payload.get("learning_objective", "")

    doc = Document()
    title = doc.add_heading("Reading Comprehension Activity", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Topic: {topic}  |  Grade: {grade}  |  Objective: {objective}")
    doc.add_paragraph("Name: ____________________________   Date: _______________")
    doc.add_paragraph()

    byr = comp.get("before_you_read", {})
    if byr:
        doc.add_heading(byr.get("title", "Before You Read"), 1)
        doc.add_paragraph(byr.get("instructions", ""))
        for i, q in enumerate(byr.get("questions", []), 1):
            num = q.get("number", i)
            doc.add_paragraph(f"{num}. {q.get('question', '')}")
            doc.add_paragraph("   Answer: ____________________________________________")
        doc.add_paragraph()

    ag = comp.get("annotation_guide", {})
    if ag:
        doc.add_heading(ag.get("title", "Annotation Guide"), 1)
        doc.add_paragraph(ag.get("instructions", ""))
        for s in ag.get("symbols", []):
            doc.add_paragraph(f"  {s.get('symbol', '')} = {s.get('meaning', '')}")
        doc.add_paragraph()

    passage = comp.get("passage", {})
    if passage:
        doc.add_heading(passage.get("title", "Reading Passage"), 1)
        for para in passage.get("text", "").split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()

    tdq = comp.get("text_dependent_questions", {})
    if tdq:
        doc.add_heading(tdq.get("title", "Text-Dependent Questions"), 1)
        doc.add_paragraph(tdq.get("instructions", ""))
        for i, q in enumerate(tdq.get("questions", []), 1):
            num = q.get("number", i)
            doc.add_paragraph(f"{num}. {q.get('question', '')}")
            doc.add_paragraph("   Answer: ____________________________________________")
            doc.add_paragraph("   ____________________________________________________")
        doc.add_paragraph()

    vic = comp.get("vocabulary_in_context", {})
    if vic:
        doc.add_heading(vic.get("title", "Vocabulary in Context"), 1)
        doc.add_paragraph(vic.get("instructions", ""))
        for i, item in enumerate(vic.get("items", []), 1):
            doc.add_paragraph(f"{i}. Word: \"{item.get('word', '')}\"")
            doc.add_paragraph(f"   From the text: \"{item.get('sentence_from_passage', '')}\"")
            doc.add_paragraph(f"   {item.get('activity', '')}")
            doc.add_paragraph("   My answer: _________________________________________")
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return _SR(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="reading_{topic}.docx"'},
    )


# ── Reading RAG document upload ──

@app.post("/api/reading/rag/add-text")
async def reading_add_rag_text(req: ReadingRAGDocRequest):
    doc_id = reading_save_rag_document(req.content, "knowledge", req.topic, req.grade_level)
    return {"success": True, "doc_id": doc_id}

@app.post("/api/reading/rag/add-file")
async def reading_add_rag_file(request: Request, file: UploadFile = File(...)):
    await upload_limiter.check(client_ip(request))
    raw = await read_upload_capped(file)
    content = ""
    if file.filename.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(raw))
        content = " ".join(p.extract_text() or "" for p in reader.pages)
    elif file.filename.endswith(".docx"):
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(raw))
        content = " ".join(p.text for p in doc.paragraphs)
    else:
        content = raw.decode("utf-8", errors="ignore")
    content = (content or "").strip()
    doc_id = reading_save_rag_document(content[:6000], "file", file.filename, 0)
    return {
        "success": True,
        "doc_id": doc_id,
        "chars_indexed": len(content),
        "text": content[:8000],
        "filename": file.filename,
    }


# ── Reading URL extraction ──

@app.post("/api/reading/extract-url")
async def reading_extract_url(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    url = assert_public_url((req.get("url") or "").strip())
    try:
        import httpx
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=20.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ReadingTool/1.0)"}) as cx:
            r = await cx.get(url)
            assert_public_url(str(r.url))
            r.raise_for_status()
            soup = BeautifulSoup(r.text, "html.parser")
            for bad in soup(["script", "style", "noscript", "iframe", "nav", "footer", "header", "form", "aside"]):
                bad.decompose()
            title = (soup.title.string or "").strip() if soup.title else ""
            main = soup.find("main") or soup.find("article") or soup.body or soup
            text = _re_mod.sub(r"\s+\n", "\n", main.get_text("\n", strip=True))
            text = _re_mod.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Could not extract readable text from this page.")
        return {"success": True, "title": title, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"URL fetch failed: {e}")


# ── Reading YouTube extraction ──

async def _reading_youtube_metadata_fallback(video_id: str, url: str) -> dict | None:
    import httpx
    title = ""
    author = ""
    description = ""
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ReadingTool/1.0)"}) as cx:
            r = await cx.get("https://www.youtube.com/oembed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}", "format": "json"})
            if r.status_code == 200:
                j = r.json()
                title = (j.get("title") or "").strip()
                author = (j.get("author_name") or "").strip()
    except Exception:
        pass
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ReadingTool/1.0)"}) as cx:
            r = await cx.get("https://noembed.com/embed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}"})
            if r.status_code == 200:
                j = r.json()
                if not title:  title  = (j.get("title") or "").strip()
                if not author: author = (j.get("author_name") or "").strip()
                description = (j.get("description") or "").strip()
    except Exception:
        pass
    try:
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                              "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept-Language": "en-US,en;q=0.9",
                "Cookie": "CONSENT=YES+cb.20210328-17-p0.en+FX+000",
            }) as cx:
            r = await cx.get(f"https://www.youtube.com/watch?v={video_id}")
            if r.status_code == 200:
                soup = BeautifulSoup(r.text, "html.parser")
                if not title:
                    ot = soup.find("meta", attrs={"property": "og:title"})
                    if ot: title = (ot.get("content") or "").strip()
                ot_desc = soup.find("meta", attrs={"property": "og:description"})
                if ot_desc and not description:
                    description = (ot_desc.get("content") or "").strip()
                for s in soup.find_all("script"):
                    txt = s.string or ""
                    if "shortDescription" in txt:
                        mm = _re_mod.search(r'"shortDescription":"((?:\\.|[^"\\])*)"', txt)
                        if mm:
                            full = mm.group(1).encode("utf-8").decode("unicode_escape", errors="ignore")
                            if len(full) > len(description):
                                description = full
                            break
    except Exception:
        pass
    pieces = []
    if title:       pieces.append(f"Video title: {title}")
    if author:      pieces.append(f"Channel: {author}")
    if description: pieces.append(f"\n{description}")
    text = "\n".join(pieces).strip()
    if not text or len(text) < 20:
        return None
    return {"title": title or f"YouTube video {video_id}", "text": text}

@app.post("/api/reading/extract-youtube")
async def reading_extract_youtube(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    url = (req.get("url") or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="url is required")
    m = _re_mod.search(r"(?:v=|youtu\.be/|/embed/|/shorts/)([A-Za-z0-9_-]{11})", url)
    if not m:
        raise HTTPException(status_code=400, detail="Could not detect a YouTube video id in that URL.")
    video_id = m.group(1)
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        proxy_config = None
        wh_user = os.getenv("WEBSHARE_PROXY_USERNAME")
        wh_pass = os.getenv("WEBSHARE_PROXY_PASSWORD")
        if wh_user and wh_pass:
            try:
                from youtube_transcript_api.proxies import WebshareProxyConfig
                proxy_config = WebshareProxyConfig(proxy_username=wh_user, proxy_password=wh_pass)
            except Exception:
                pass
        if proxy_config is not None:
            transcript = YouTubeTranscriptApi(proxy_config=proxy_config).fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        elif hasattr(YouTubeTranscriptApi, "get_transcript"):
            chunks = YouTubeTranscriptApi.get_transcript(video_id)
            text = " ".join((c.get("text", "") if isinstance(c, dict) else getattr(c, "text", "")) for c in chunks).strip()
        else:
            transcript = YouTubeTranscriptApi().fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Transcript was empty.")
        return {"success": True, "video_id": video_id, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except Exception as e:
        msg = str(e)
        blocked = ("blocking requests" in msg.lower()
                   or ("ip" in msg.lower() and "block" in msg.lower())
                   or "could not retrieve a transcript" in msg.lower())
        if blocked:
            fallback = await _reading_youtube_metadata_fallback(video_id, url)
            if fallback:
                return {
                    "success": True, "video_id": video_id, "url": url,
                    "text": fallback["text"][:8000], "chars": len(fallback["text"]),
                    "title": fallback["title"],
                    "note": "Transcript was blocked -- using title + description instead.",
                }
        raise HTTPException(
            status_code=502 if blocked else 500,
            detail=(
                "YouTube blocks transcript requests from our server. Please paste the transcript manually."
                if blocked else f"YouTube transcript fetch failed: {msg}"
            ),
        )


# ── Reading Auto-fields ──

@app.post("/api/reading/auto-fields")
async def reading_auto_fields(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    source_text = (req.get("source_text") or "").strip()
    if not source_text:
        raise HTTPException(status_code=400, detail="source_text is required.")
    grade = req.get("grade_level")
    grade_hint = f"Calibrate the topic + objective for Grade {grade} students. " if grade else ""
    prompt = (
        "You are an expert reading-comprehension curriculum designer.\n"
        "Read the SOURCE MATERIAL below and propose a clean, classroom-ready\n"
        "  * Topic (a short noun phrase, 4-10 words, naming the main idea)\n"
        "  * Learning Objective (one sentence starting with 'Students will...' "
        "describing what the student will be able to do after reading).\n"
        f"{grade_hint}"
        "Return ONLY a JSON object with keys 'topic' and 'learning_objective'. "
        "No markdown, no prose outside the JSON.\n\n"
        "SOURCE MATERIAL:\n---\n"
        f"{source_text[:6000]}\n---\n\n"
        'JSON: {"topic": "...", "learning_objective": "Students will ..."}'
    )
    try:
        completion = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You return strict JSON. No markdown fences."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=400,
            response_format={"type": "json_object"},
        )
        raw = completion.choices[0].message.content.strip()
        data = json.loads(raw)
        topic = (data.get("topic") or "").strip()
        objective = (data.get("learning_objective") or "").strip()
        if not topic and not objective:
            raise HTTPException(status_code=502, detail="Model returned no fields.")
        return {"success": True, "topic": topic, "learning_objective": objective}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"auto-fields failed: {e}")


# ─── SERVE FRONTEND ────────────────────────────────────

FRONTEND_DIST = Path(__file__).parent.parent / "frontend" / "dist"

if FRONTEND_DIST.exists():
    app.mount("/assets", StaticFiles(directory=FRONTEND_DIST / "assets"), name="assets")

NO_CACHE_HEADERS = {"Cache-Control": "no-cache, no-store, must-revalidate", "Pragma": "no-cache", "Expires": "0"}

@app.get("/")
def serve_index():
    index_file = FRONTEND_DIST / "index.html"
    if index_file.exists():
        return FileResponse(index_file, headers=NO_CACHE_HEADERS)
    return {"message": "Frontend not built"}

@app.get("/{full_path:path}", include_in_schema=False)
def serve_frontend(full_path: str):
    if full_path.startswith("api/"):
        return {"error": "Not found"}
    file_path = FRONTEND_DIST / full_path
    if file_path.exists():
        return FileResponse(file_path)
    index_file = FRONTEND_DIST / "index.html"
    if index_file.exists():
        return FileResponse(index_file, headers=NO_CACHE_HEADERS)
    return {"error": "Not found"}

if __name__ == "__main__":
    import uvicorn
    import socket

    # Configure socket to allow reuse (fixes TIME_WAIT port binding issues on Windows)
    class SocketReusableServer(uvicorn.Server):
        def install_signal_handlers(self):
            pass

    # Create uvicorn config with socket reuse enabled
    config = uvicorn.Config(
        app,
        host="127.0.0.1",
        port=8001,
        log_level="info"
    )

    # Allow address reuse to fix "port already in use" errors
    config.disable_lifespan = False

    server = uvicorn.Server(config)

    # Patch socket to allow reuse
    original_socket = socket.socket
    def patched_socket(*args, **kwargs):
        sock = original_socket(*args, **kwargs)
        sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        if hasattr(socket, 'SO_REUSEPORT'):
            sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEPORT, 1)
        return sock
    socket.socket = patched_socket

    try:
        import asyncio
        asyncio.run(server.serve())
    except KeyboardInterrupt:
        pass
